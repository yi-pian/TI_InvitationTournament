"""Targeted revision of the two existing manuals.

This script loads the DOCX files in place.  It does not rebuild either manual
from the old generator: only the reviewed blocks are replaced or inserted.
"""
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

FULL = '电子设计竞赛单片机_信号题_完整Debug手册.docx'
QUICK = '电子设计竞赛单片机_信号题_比赛现场速查手册.docx'
BLUE = '1F4E78'

P2_CARDS = {
    '15': ('VGA / PGA / AD603 增益链', '需要外置可变增益前端时；当前工程未集成。', '同时测输入、增益控制端、输出与 ADC raw min/max。', '增益控制电压、供电或 ADC 余量不清时，STOP：不要采用测幅结果。', '07_signal_frontend 的前端资料；先固定单档。', '固定中档、显示 OVER RANGE，暂停自动量程。'),
    '16': ('比较器 / LM393 / 零交叉', '需要外置比较器或零交叉整形时；当前工程未集成。', '示波器同时看比较器输入、阈值和输出。', '输出边沿或上拉状态未证实时，STOP：不要查 Timer 捕获算法。', 'comparator_zero_cross/README.md。', '改用 ADC 软件零交叉并加滞回。'),
    '18': ('信号完整性', '高频、长线或异常串扰时使用；当前工程无专用 SI 模块。', '用短地弹簧比较信号源端和 ADC 引脚。', '探头/接地方式未证实时，STOP：不要把观测伪影当软件问题。', '板级 SI 原则与实际原理图。', '降频、缩线、只保留数值显示。'),
    '19': ('模拟滤波', '需要抗混叠或限制测量带宽时；当前工程未集成可配置滤波。', '测滤波器前后幅度、相位和截止附近响应。', '截止频率与 Fs 关系不清时，STOP：不要拿滤波后数据做精密结果。', '模拟前端资料。', '提高 Fs 或降低测量带宽。'),
    '20': ('数字滤波', '原始帧正确但需要降噪时；当前工程未集成滤波模块。', '比较滤波前后结果、延迟和 CPU 时间。', 'raw 或帧所有权未验证时，STOP：不引入复杂滤波。', '算法目录中的最小例程，先离线验证。', '关闭滤波或用短窗口平均。'),
    '22': ('I2C', '接入 I2C 传感器/外设时；当前工程未集成。', '空闲时测 SCL/SDA 是否释放为高，再看首个 ACK。', 'SCL/SDA 未释放为高时，STOP：不要调驱动状态机。', '目标器件的最小读写例程。', '取消外设，固定参数。'),
    '23': ('UART', '需要串口日志/通信时；当前工程未集成。', '发送短固定文本并测 TX 位宽，确认端口与波特率。', '串口电平或波特率未证实时，STOP：不要用日志结论判断算法。', '平台最小 UART 例程。', '屏幕显示状态码或 LED。'),
    '25': ('按键与键盘', '需要新增/迁移键盘接口时；当前工程已有键盘资源，通用逻辑不另加 API。', '先看 GPIO 电平，再验证单次按下是否只触发一次。', 'GPIO 电平本身不可靠时，STOP：不要改 UI 状态机。', '当前工程键盘配置及 GPIO 最小例程。', '固定参数，暂停交互功能。'),
    '28': ('FPGA + MCU 协同', '题目采用 FPGA 接口时；当前工程未集成。', '逻辑分析仪看时钟、片选、数据、有效标志和帧边界。', '时钟域/电平/协议未单独验证时，STOP：不要连完整算法链。', 'FPGA 侧时序图和最小通信例程。', 'MCU 独立采集/计算，FPGA 仅保留可选加速。'),
    '30': ('精度增强', '基础结果稳定但误差仍不够时；当前工程未集成精度算法。', '记录真值、测量值、Fs、量程和误差趋势。', '基线精度与 Fs 未验证时，STOP：不要先上插值/校正。', '05_precision 目录。', '使用未插值 FFT 或零交叉结果。'),
    '31': ('猝发与触发', '短脉冲/预触发采集需求；当前工程未集成。', '同步观察触发输入、采样节拍和 DMA 完成。', '触发源/时间基准不清时，STOP：不要加复杂缓存。', 'ADC/DMA 环形缓冲例程。', '连续采样，人工保存最近一帧。'),
}

def text_of(el):
    return ''.join(node.text or '' for node in el.iter() if node.tag == qn('w:t'))

def style_of_p(el):
    pPr = el.find(qn('w:pPr'))
    if pPr is None:
        return ''
    st = pPr.find(qn('w:pStyle'))
    return '' if st is None else (st.get(qn('w:val')) or '')

def body_children(doc):
    return list(doc.element.body)

def is_h1(el):
    return el.tag == qn('w:p') and style_of_p(el) in ('Heading1', 'Heading 1')

def find_debug_block(doc, prefix):
    children = body_children(doc)
    start = next(i for i, el in enumerate(children) if el.tag == qn('w:p') and text_of(el).startswith(prefix))
    end = len(children) - 1
    for i in range(start + 1, len(children)):
        if is_h1(children[i]):
            end = i
            break
    return start, end

def remove_preceding_pagebreak(doc, start):
    children = body_children(doc)
    if start == 0:
        return
    prev = children[start - 1]
    if prev.tag == qn('w:p') and 'w:type="page"' in prev.xml:
        doc.element.body.remove(prev)

def basic_table(doc, headers, rows, size=8.4):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    trPr = t.rows[0]._tr.get_or_add_trPr(); hdr = OxmlElement('w:tblHeader'); hdr.set(qn('w:val'), 'true'); trPr.append(hdr)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = h
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), BLUE); cell._tc.get_or_add_tcPr().append(shd)
        for r in cell.paragraphs[0].runs:
            r.bold = True; r.font.color.rgb = None; r.font.size = Pt(size)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                for r in p.runs: r.font.size = Pt(size)
    return t

def fragment_children(temp):
    return [deepcopy(x) for x in list(temp.element.body) if x.tag != qn('w:sectPr')]

def replace_block(doc, prefix, temp, remove_pagebreak=False):
    start, end = find_debug_block(doc, prefix)
    if remove_pagebreak:
        remove_preceding_pagebreak(doc, start)
        start, end = find_debug_block(doc, prefix)
    body = doc.element.body
    children = body_children(doc)
    anchor = children[start]
    for child in children[start + 1:end]:
        body.remove(child)
    pos = list(body).index(anchor)
    body.remove(anchor)
    for i, child in enumerate(fragment_children(temp)):
        body.insert(pos + i, child)

def add_section(temp, title, text):
    temp.add_paragraph(title, style='Heading 2')
    temp.add_paragraph(text)

def dds_fragment():
    d = Document()
    d.add_paragraph('Debug 14-A｜软件 DDS', style='Heading 1')
    basic_table(d, ['优先级', '工程状态', '定位'], [('P1', '【当前工程未实现】', 'MCU Timer/DMA → 相位累加器 → 查表 → DAC → 重建滤波')])
    add_section(d, '现场现象', '固定频率无输出、频率按比例错误、扫频失真或界面一开便卡顿。')
    add_section(d, 'STOP 条件', 'DAC 静态三档电压、Timer/DMA 实际更新率或 DAC 引脚波形未证实前，STOP：不能进入扫频、调制或频谱讨论。')
    add_section(d, '数据链', 'Timer/DMA 以实际 Fupdate 取样并送出波形表；相位累加器按 FTW 前进；DAC 输出经重建低通滤波后形成模拟波形。')
    add_section(d, '第一测量', '依次测 Timer/DMA 更新节拍、DAC 引脚、重建滤波器后输出；打印 Fupdate、FTW、波形表长度。')
    add_section(d, '核心公式', 'fout = FTW × Fupdate / 2^N。其中 N 是相位累加器位宽，Fupdate 是 DAC 的实际更新率，不是 CPU 主频。')
    add_section(d, '正常参考值', 'Fupdate 必须由运行参数或硬件节拍确认；输出频率随 FTW 单调变化；滤波器后波形无明显阶梯镜像。')
    add_section(d, '典型错误值 / 异常分支', '频率比例错：Fupdate 错；无输出：Timer/DMA/DAC 未起；失真：表长或更新率不足、DAC 削顶、重建滤波不合适。')
    add_section(d, '工程 / 代码 / 硬件检查', '【当前工程未实现】不得假定已有 DDS API。检查相位累加器溢出逻辑、查表索引、DMA 目的与 DAC 触发；检查 DAC 参考、输出负载和重建滤波。')
    add_section(d, '最小 example', '先使用项目 06_generator/dds/README_MINIMAL_EXAMPLE.c 的单频路径；固定 1 kHz 后才验证多频点和扫频。')
    add_section(d, '修复后验收', '固定 1 kHz、10 kHz 等频点稳定，按实测 Fupdate 计算后误差可解释；接入滤波后波形与频率均不漂移。')
    add_section(d, '保底方案', '手动固定频点；必要时改用外部信号源。')
    basic_table(d, ['30 秒 Checklist', '结论'], [('Timer/DMA 节拍、DAC、滤波后三点是否依次成立？', '任一前级未成立即 STOP，不查扫频算法。')], 8)
    d.add_paragraph('Debug 14-B｜外置 DDS：AD9833 / AD9850', style='Heading 1')
    basic_table(d, ['优先级', '工程状态', '定位'], [('P1', '【当前工程未实现】', '外置 DDS 由其参考时钟与专用装载接口决定；不能套用软件 DDS 的 Fupdate。')])
    add_section(d, '现场现象', 'MCU 有通信但无输出、输出频率错、复位后失效，或更改频率字无响应。')
    add_section(d, 'STOP 条件', '参考时钟、复位和芯片输出未分别证实时，STOP：不要先修改频率字算法或上扫频。')
    add_section(d, 'AD9833：公式与测点', 'FREQREG = round(fout × 2^28 / MCLK)。先测 MCLK，再测 FSYNC、SCLK、SDATA、RESET，最后测正弦/方波输出。AD9833 使用其三线串行接口。')
    add_section(d, 'AD9850：公式与测点', 'FTW = round(fout × 2^32 / CLKIN)。依次测 CLKIN、RESET、W_CLK、FQ_UD、DATA 和输出。AD9850 不是普通 SPI：它使用自身的串行装载或并行字节装载接口。')
    add_section(d, '第一测量', '先确认参考时钟幅度/频率与 RESET 释放，再确认装载时序；最后分别在芯片模拟输出和模块后级输出测量。')
    add_section(d, '正常参考值', 'AD9833 的频率寄存器宽度为 28 bit；AD9850 的调谐字宽度为 32 bit。FTW/FREQREG 必须按各自参考时钟计算。')
    add_section(d, '典型错误值 / 异常分支', '频率比例错：MCLK/CLKIN 参数错；恒定无输出：RESET 未释放、参考时钟/供电错误；通信无效：把 AD9850 当普通 SPI，或 AD9833 的 FSYNC 时序错误。')
    add_section(d, '工程 / 代码 / 硬件检查', '【当前工程未实现】从所用芯片数据手册和模块原理图确认接口、电平与装载次序；不得新增未经验证的驱动 API。')
    add_section(d, '最小 example', '一次只输出固定 1 kHz；AD9833 与 AD9850 分别建立最小驱动，不共享频率字公式和时序。')
    add_section(d, '修复后验收', '在至少三个固定频点下，参考时钟、控制时序与输出频率一致；复位后可重复装载。')
    add_section(d, '保底方案', '固定已验证频点，或用外部信号源替代。')
    basic_table(d, ['30 秒 Checklist', '结论'], [('先识别芯片，再查其参考时钟、RESET、专用接口、输出。', 'AD9833 与 AD9850 的公式和接口绝不混用。')], 8)
    return d

def speed_fragment(number, spec):
    title, use, first, stop, example, fallback = spec
    d = Document()
    d.add_paragraph('Debug %s｜%s（通用模块速卡）' % (number, title), style='Heading 1')
    basic_table(d, ['适用条件', '第一测量', 'STOP', '最小例程', '保底方案'], [(use, first, stop, example, fallback)], 8)
    return d

def replace_all(doc, old, new):
    def visit(paras):
        for p in paras:
            if old in p.text:
                p.text = p.text.replace(old, new)
    visit(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                visit(cell.paragraphs)
                for nt in cell.tables:
                    for nr in nt.rows:
                        for nc in nr.cells: visit(nc.paragraphs)

def update_resource_matrix(doc):
    for table in doc.tables:
        for row in table.rows:
            texts = [c.text.strip() for c in row.cells]
            if texts and texts[0] == 'ADC1':
                row.cells[1].text = 'PA17 / ADC1 MEM0 ch.2'
                row.cells[2].text = '同步采样 Y 路，用于当前相位/Lissajous 流程；前端信号含义以题目硬件连接为准，DMA CH1。'
            if texts and texts[0] == 'SPI':
                row.cells[1].text = 'SPI1：PB9 SCK / PB8 PICO(MOSI) / PB6 CS0 / PA16 POCI(MISO)'
                row.cells[2].text = 'ILI9341 总线；PA16 已由当前 SysConfig 分配，勿随意复用。'
            if texts and texts[0] == 'TFT GPIO':
                row.cells[1].text = 'PB15 DC / PB12 BLK'
                row.cells[2].text = 'ILI9341 控制脚；SPI 数据脚见上一行。'

def insert_after(parent, anchor, items):
    pos = list(parent).index(anchor) + 1
    for i, item in enumerate(items): parent.insert(pos + i, item)

def source_truth_fragment():
    d = Document()
    d.add_paragraph('采样率唯一事实源', style='Heading 1')
    basic_table(d, ['项目', '当前工程事实', '禁止误用'], [
        ('当前运行目标', 'main.c：SIGNAL_SAMPLE_RATE_HZ = 500000U。', '不要把 signal_config.h 的 100 kSPS 模板值当运行值。'),
        ('实际运行值', 'SignalDualADC_GetConfiguredRate()。', '不要只根据生成的 Timer Load 判断实际 Fs。'),
        ('配置关系', 'SignalDualADC_SetSampleRate() 在运行时按 Timer Clock 重写 Timer Load。SysConfig 参数仅作初始化模板。', '生成文件与 .syscfg 用于资源/初始化核对，不替代运行时配置。'),
    ], 8.2)
    return fragment_children(d)

def insert_source_truth(doc, title_prefix):
    children = body_children(doc)
    heading_idx = next(i for i,e in enumerate(children) if e.tag == qn('w:p') and text_of(e).startswith(title_prefix))
    anchor = next(e for e in children[heading_idx+1:] if e.tag == qn('w:tbl'))
    insert_after(doc.element.body, anchor, source_truth_fragment())

def append_preflight_and_prohibitions(doc, compact=False):
    d = Document()
    d.add_paragraph('比赛前 30 分钟检查表', style='Heading 1')
    basic_table(d, ['类别', '必须确认'], [
        ('硬件', '电源与模块供电稳定；烧录线可用；示波器探头已补偿；信号源输出正确；接线/共地/极性复核。'),
        ('软件', '工程备份可打开；Clean Build 成功；连续烧录成功；比赛参数已冻结；完成 Git 提交或压缩备份。'),
    ], 8.2)
    d.add_paragraph('禁止操作清单', style='Heading 1')
    basic_table(d, ['出现条件', '禁止', '立即做'], [
        ('最后 30 分钟', '禁止改 SysConfig、一次改多个模块或大规模重构。', '仅修明确 Bug；参数冻结；保留可烧录基线。'),
        ('ADC raw 错', '禁止调 FFT。', '回到 Debug 06，看 g_raw_a/g_raw_b 与 ADC 引脚。'),
        ('DMA 未完成', '禁止跑算法。', '回到 Debug 07，只消费完成帧。'),
        ('供电异常', '禁止查通信。', '断扩展模块，测 3.3 V / 电流 / 复位。'),
    ], 8.1)
    items = fragment_children(d)
    # append before sectPr
    body = doc.element.body
    pos = len(list(body)) - 1
    for i, item in enumerate(items): body.insert(pos + i, item)

def remove_all_hard_pagebreaks(doc):
    body = doc.element.body
    for el in list(body):
        if el.tag == qn('w:p') and 'w:type="page"' in el.xml:
            body.remove(el)

def revise_full():
    doc = Document(FULL)
    update_resource_matrix(doc)
    insert_source_truth(doc, '第 4 章')
    replace_all(doc, '第一测量：在 Timer ISR/触发点翻转 GPIO，用示波器测周期；再由相邻采样点或捕获验证。',
                '第一测量：优先观测硬件 Timer 输出/事件节拍；其次读取 SignalDualADC_GetConfiguredRate()、Timer Clock 与实际 Timer Load；最后才用临时 GPIO 诊断。GPIO 翻转只能短时、低侵入诊断，不能长期加入 500 kSPS 实时采样链。')
    replace_all(doc, '先记录 baseline，再用 Δt/相位校正；只对稳定、不过零歧义的信号报告相位。',
                '同源 baseline 只适用于 frequency_ratio=1 的同频模式；此时可用 Δt = Δφ/(360f)。对 2~5 倍频模式，必须在对应倍率下建立独立 baseline；不满足条件时只报告 INVALID。')
    replace_all(doc, 'Debug 24｜ILI9341、OLED 与屏幕', 'Debug 24｜ILI9341显示')
    replace_all(doc, '例：Fs=5 MHz、N=4096，则 Δf≈1220.703 Hz；100 kHz 对应 k≈81.92，峰应在 82 附近。',
                '【通用算法示例，不代表当前工程配置】Fs=5 MHz、N=4096，则 Δf≈1220.703 Hz；100 kHz 对应 k≈81.92，峰应在 82 附近。当前工程为 Fs=500 kSPS、N=1024，Δf=488.28125 Hz，且 FFT 尚未集成。')
    replace_all(doc, '【当前工程未集成 FFT】先核对 fft_minimal.c 的接口与内存需求，禁止假设工程已有 FFT API。',
                '【当前工程未集成 FFT】先核对 fft_minimal.c 的接口与内存需求；上面的 5 MHz/4096 是通用算例，不可当作当前工程参数，禁止假设工程已有 FFT API。')
    replace_block(doc, 'Debug 14｜', dds_fragment())
    for number, spec in P2_CARDS.items():
        replace_block(doc, 'Debug %s｜' % number, speed_fragment(number, spec), remove_pagebreak=True)
    # OLED is intentionally not presented as a current-engineering display path.
    d = Document(); d.add_paragraph('未集成通用模块说明', style='Heading 1'); d.add_paragraph('OLED：当前工程未实现 OLED 驱动或资源分配。若比赛题目改用 OLED，应以器件模块最小例程重新建立资源表；不得把 ILI9341 的 SPI/GPIO 初始化直接套用。')
    body = doc.element.body; pos = len(list(body)) - 1
    for i, item in enumerate(fragment_children(d)): body.insert(pos + i, item)
    append_preflight_and_prohibitions(doc)
    doc.save(FULL)

def quick_dds_fragment():
    d = Document()
    d.add_paragraph('Debug 14-A｜软件 DDS', style='Heading 2')
    basic_table(d, ['第一测量', 'STOP', '公式 / 保底'], [('测 Timer/DMA 更新率、DAC 引脚、重建滤波后输出。', 'DAC 静态输出或 Fupdate 未证实，不进入扫频。', 'fout=FTW×Fupdate/2^N；固定频点或外部信号源。')], 7.8)
    d.add_paragraph('Debug 14-B｜外置 DDS', style='Heading 2')
    basic_table(d, ['第一测量', 'STOP', '公式 / 保底'], [('AD9833：MCLK/FSYNC/SCLK/SDATA/RESET；AD9850：CLKIN/RESET/W_CLK/FQ_UD/DATA。', '参考时钟、RESET、专用接口未证实，不改频率字。', 'AD9833：round(fout×2^28/MCLK)；AD9850：round(fout×2^32/CLKIN)。AD9850 非普通 SPI；固定频点。')], 7.6)
    return d

def revise_quick():
    doc = Document(QUICK)
    update_resource_matrix(doc)
    insert_source_truth(doc, '第 4 页')
    replace_all(doc, '第一测量：在 Timer ISR/触发点翻转 GPIO，用示波器测周期；再由相邻采样点或捕获验证。',
                '优先硬件 Timer 输出/事件；其次读 GetConfiguredRate、Timer Clock、Timer Load；最后才临时 GPIO，且不得长期进入 500 kSPS 采样链。')
    replace_all(doc, '先记录 baseline，再用 Δt/相位校正；只对稳定、不过零歧义的信号报告相位。',
                '仅 frequency_ratio=1 同频时用同源 baseline 与 Δt=Δφ/(360f)；2~5 倍频须建立对应倍率 baseline。')
    replace_all(doc, 'Debug 24｜ILI9341、OLED 与屏幕', 'Debug 24｜ILI9341显示')
    replace_block(doc, 'Debug 14｜', quick_dds_fragment())
    # The old fast manual used sixteen fixed breaks. Let Word flow compact cards naturally.
    remove_all_hard_pagebreaks(doc)
    append_preflight_and_prohibitions(doc, compact=True)
    # Clear generic DDS formula that mixed the two DDS classes; their two mini-cards are authoritative.
    replace_all(doc, 'FTW≈fout×2^N/Fupdate', '软件 DDS：fout=FTW×Fupdate/2^N；外置 DDS 见 Debug 14-B')
    doc.save(QUICK)

if __name__ == '__main__':
    revise_full()
    revise_quick()
    print('revised:', FULL)
    print('revised:', QUICK)
