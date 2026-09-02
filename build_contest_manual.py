from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path('电子设计竞赛单片机_信号题_比赛现场问题与解决方法大全.docx')

BLUE = '2E74B5'; DARK = '1F4D78'; PALE = 'E8EEF5'; NOTE = 'F4F6F9'; RED = '9B1C1C'

def fonts(run, size=None, color=None, bold=None):
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def borders(cell):
    tcPr = cell._tc.get_or_add_tcPr(); mar = OxmlElement('w:tcMar')
    for side in ('top','start','bottom','end'):
        e = OxmlElement('w:' + side); e.set(qn('w:w'), '80' if side in ('top','bottom') else '120'); e.set(qn('w:type'), 'dxa'); mar.append(e)
    tcPr.append(mar)

def set_table_widths(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
    ind = OxmlElement('w:tblInd'); ind.set(qn('w:w'), '120'); ind.set(qn('w:type'), 'dxa'); tblPr.append(ind)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn('w:tcW'))
            if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'), 'dxa'); borders(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    trPr = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader'); header.set(qn('w:val'), 'true'); trPr.append(header)

def set_cell(cell, text, bold=False, color=None, size=9.5):
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text); fonts(r, size, color, bold)

def add_text(doc, text, bold_lead=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead); fonts(r, 10.5, DARK, True); r = p.add_run(text[len(bold_lead):]); fonts(r, 10.5)
    else:
        r = p.add_run(text); fonts(r, 10.5)
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.18
        fonts(p.add_run(item), 10.2)

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}'); p.paragraph_format.keep_with_next = True
    r = p.add_run(text); fonts(r, {1:16,2:13,3:11.5}[level], BLUE if level < 3 else DARK, True)
    return p

def add_callout(doc, title, text):
    t = doc.add_table(rows=1, cols=1); set_table_widths(t,[9360]); c=t.cell(0,0); shade(c, NOTE)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(2); r=p.add_run(title+'：'); fonts(r,10.5,DARK,True); r=p.add_run(text); fonts(r,10.5)

CARDS = [
('1. 整机完全不工作','【比赛高频】【致命】','P0','发生概率：★★★★☆  影响程度：★★★★★  排查难度：★★☆☆☆','上电后无灯、限流、电压反复掉下去或接入模块即重启。','短路/反接；DC-DC 限流或启动浪涌；5 V/3.3 V 用错；正负电源缺失；未共地。','万用表先量输入、3.3 V、5 V、运放正负电源；断开外设逐个回接；看电源电流而不是只看电压。','1) 断开所有扩展板；2) 空载确认各轨；3) 查电阻和极性；4) 按“MCU→ADC→屏→DDS”逐件接回；5) 最后接被测系统地。','短路先清除；增大供电余量和去耦；屏幕/DDS 单独加大电容；所有信号源与板共地。','推荐：最小系统先稳定运行，再逐层加模块。','备用：关闭背光、降低屏幕刷新和DDS输出，保留最核心测量链。','[ ] 限流是否触发  [ ] 3.3/5 V 是否正确  [ ] 正负电源是否完整  [ ] 共地  [ ] 是否某一模块一接即掉压'),
('2. 单片机无法烧录 / 无法启动','【比赛高频】【致命】【快速修复】','P0','发生概率：★★★☆☆  影响程度：★★★★★  排查难度：★★☆☆☆','ST-Link/XDS110 找不到目标、下载报错、复位后无任何串口/LED。','调试器供电与外部供电冲突；SWD/JTAG 接反；Reset/Boot 拉错；时钟或启动代码异常；程序启动即 HardFault。','先只接调试器，观察目标电压；用空 LED 程序验证；调试暂停于 Reset_Handler；检查复位脚电平。','1) 断外设；2) 只留调试器与 GND；3) 降速下载；4) Mass erase 后下载最小程序；5) 再恢复工程。','改为可靠的调试线和共地；把启动阶段外设初始化逐项注释定位。','推荐：保留“LED+串口”救援工程。','备用：拔除可能占用调试脚的模块，重新导入干净工程。','[ ] 目标电压  [ ] SWD/XDS110 接线  [ ] Reset  [ ] Boot  [ ] 最小程序能否跑'),
('3. CCS / CubeMX / Keil / SysConfig 工程配置','【历史真实问题】【比赛高频】','P0','发生概率：★★★★☆  影响程度：★★★★★  排查难度：★★★☆☆','编译、链接或生成文件错误；导入后外设宏缺失；配置改了但行为没变。','SysConfig 未 Generate；实例名/PinMux 与模块 README 不一致；新 .c 被 Exclude from Build；include/linker 路径错；直接手改生成文件。','先看 Problems 的第一条错误；确认 .syscfg 保存并重新生成；检查 CCS Project Explorer 的编译包含文件。','1) 从第一条编译错误处理；2) Generate；3) Clean；4) 确认 .c 被编译；5) Build；6) 上板。','以 README 的实例名和点击路径为唯一依据；生成文件只核对不手改。','推荐：每加一个模块就 Build 一次。','备用：回退到空母版，逐模块复制。','[ ] syscfg 已保存/生成  [ ] 实例名匹配  [ ] .c 未排除  [ ] include 路径  [ ] 第一条错误已清'),
('4. ADC 采集问题','【比赛高频】【致命】','P0','发生概率：★★★★★  影响程度：★★★★★  排查难度：★★★☆☆','数据全 0、满量程、乱跳、失真、双通道不同步或 DMA 缓冲覆盖。','输入悬空/未共地；输入范围或偏置错；触发频率错；DMA 长度/地址错；ADC 饱和；外置 ADC 时序错。','示波器同时看 ADC 输入与触发；串口打印首/中/尾原始码；统计 min/max；查看 DMA 完成标志。','1) 固定 DC 输入；2) 正弦低频；3) 查原始码值与参考电压；4) 查定时器触发；5) 再查 DMA；6) 最后进算法。','给 ADC 输入偏置并限幅；保证采样率；DMA 缓冲和算法缓冲分离；双路必须由同一 timer/event 触发。','推荐：先单次 DMA 验证，再 Continuous/Ping-Pong。','备用：缩短缓冲并降低采样率，先保住一帧测量。','[ ] 共地  [ ] min/max  [ ] 未饱和  [ ] Fs  [ ] DMA 完成  [ ] 双路同触发'),
('5. 测频问题','【比赛高频】','P1','发生概率：★★★★☆  影响程度：★★★★☆  排查难度：★★★☆☆','高频偏大/偏小、低频刷新慢、数字跳动，方波可测而正弦不可测。','Timer 时基/分频错误；比较器门限抖动；周期数太少；零交叉受 DC 与噪声影响；FFT bin 分辨率不够。','同时打印计数值、时基、有效边沿数；示波器看比较器边沿；对比函数发生器。','1) 先确定信号范围；2) 方波用输入捕获；3) 正弦先整形；4) 增加测量窗；5) 用 FFT/插值复核。','低频用周期法或多周期平均；高频用计数法/FFT；正弦加入迟滞比较器或软件零交叉插值。','推荐：Timer 为快速显示，FFT 为高精度复核。','备用：仅显示稳定的滑动平均结果。','[ ] Timer 时钟  [ ] prescaler/ARR  [ ] 门限/迟滞  [ ] 有效周期数  [ ] 量程'),
('6. FFT / 频谱分析','【比赛高频】','P1','发生概率：★★★★☆  影响程度：★★★★☆  排查难度：★★★★☆','主峰错位、频率跳动、幅值偏低、DC 淹没频谱、THD 异常。','非相干采样导致泄漏；未去均值；窗函数/幅值修正缺失；Fs/N 换算错误；混叠；bin 搜索包含 DC。','保存一帧 ADC 原始数组；串口输出最大 bin、bin±1、Fs、N；离线用已知正弦验证。','1) 检查 ADC；2) 去 DC；3) 检查窗与 coherent gain；4) 排除 bin0；5) 限定搜索频段；6) 做三点插值。','统一用 f=k·Fs/N；需要幅值则校正窗增益；必要时做相干采样/增加 N。','推荐：Hann 窗 + 三点抛物线插值。','备用：不加窗但强制相干采样，或回退 Timer 测频。','[ ] Fs/N  [ ] 去DC  [ ] 窗  [ ] 最大bin  [ ] bin边界  [ ] 混叠'),
('7. 幅值测量','【比赛高频】','P1','发生概率：★★★★☆  影响程度：★★★★☆  排查难度：★★★☆☆','Vpp/Vrms 随噪声跳动、量程变化后不准、削顶时结果离谱。','ADC offset/gain 未校准；偶发毛刺；饱和；把 AC RMS 与总 RMS 混用；VGA 增益误差。','打印 min/max/mean/RMS、饱和样本数和增益档；示波器确认真实 Vpp。','1) 先校 ADC 码-电压；2) 检测削顶；3) 去 DC；4) 对毛刺做 Hampel/MAD；5) 再选 Vpp 或 RMS。','一般波形用分位数 Vpp；正弦用 RMS 或拟合；记录每个量程的增益与零偏。','推荐：分位数 Vpp + AC RMS 双显示。','备用：固定量程并只报不饱和的峰峰值。','[ ] ADC 标定  [ ] 饱和样本  [ ] 去DC  [ ] 毛刺  [ ] 量程/增益'),
('8. 相位测量','【历史讨论问题】【比赛高频】','P1','发生概率：★★★★☆  影响程度：★★★★☆  排查难度：★★★★☆','相位忽正忽负、跨 ±180° 跳变、改变频率后系统性偏移。','双通道非同触发；通道采样/模拟滤波延迟不同；零交叉门限不同；FFT 主 bin 不一致；未做相位展开。','对同一信号分两路输入，测“零相位基线”；打印两路 peak bin、复数相位和时间戳。','1) 同源同相基线；2) 确认同 timer 触发；3) 校准固定延时；4) 统一频率 bin；5) unwrap 后再求群时延。','固定延时用 phi_cal=360*f*dt 扣除；频率变化时用相位模型补偿。','推荐：双 ADC 同触发 + FFT 复相位。','备用：方波/正弦可用整形后的零交叉法。','[ ] 同触发  [ ] 同源基线  [ ] 延时校准  [ ] bin一致  [ ] unwrap'),
('9. 波形识别','【推导问题】','P2','发生概率：★★★☆☆  影响程度：★★★☆☆  排查难度：★★★☆☆','正弦、方波、三角、脉冲在噪声或幅值变化时误判。','只用单一阈值；未归一化；采样周期不足；带宽/滤波改变边沿；削顶。','显示归一化波形和特征值（crest factor、斜率、占空比、谐波比）。','1) 先判是否周期；2) 去 DC/归一化；3) 提取时域特征；4) FFT 谐波复核；5) 低置信度输出“未知”。','规则分类优先于复杂模型：正弦低 THD、方波高奇次谐波、三角谐波衰减快、脉冲占空比低。','推荐：时域特征 + 谐波比双判据。','备用：显示“待确认”并保留原始波形。','[ ] 周期数  [ ] 归一化  [ ] 饱和  [ ] 特征阈值  [ ] 未知类'),
('10. 猝发 / 瞬态 / 非连续信号','【推导问题】【比赛高频】','P1','发生概率：★★★☆☆  影响程度：★★★★☆  排查难度：★★★★☆','只抓到信号尾部、触发漏检、屏幕看到的波形不完整。','软件轮询过慢；无预触发；门限太高/太低；环形缓冲覆盖；DMA block 与算法争用。','输出触发时间、前后样本数、环形序号；示波器同时看比较器触发与 ADC trigger。','1) 确定门限和迟滞；2) 环形 DMA 常开；3) 锁定预触发块；4) 采后触发样本；5) 冻结快照处理。','比较器硬触发优先，软件阈值为备；预触发至少覆盖一个预期周期。','推荐：DMA 环形缓冲 + 硬件触发。','备用：降低 Fs 并延长采集窗口。','[ ] 触发门限  [ ] 迟滞  [ ] 预触发  [ ] 后触发  [ ] 缓冲未覆盖'),
('11. DAC 问题','【推导问题】','P1','发生概率：★★★☆☆  影响程度：★★★★☆  排查难度：★★★☆☆','无输出、幅度不对、阶梯明显、更新后波形失真。','初始化/使能遗漏；码值换算或参考电压错；更新率不足；负载太重；DMA/Timer 未触发。','示波器先看 DC 中点，再看固定码和低频正弦；打印 DAC code 与更新率。','1) 固定中码；2) 固定高/低码；3) 检查参考和输出范围；4) 查 timer/DMA；5) 再上查表。','输出加缓冲；Fs 至少显著高于最高频率；用 DMA 减少抖动。','推荐：固定码→低频→DMA 查表三步验证。','备用：降低输出频率或改 DDS 模块。','[ ] 使能  [ ] Vref  [ ] 码值  [ ] 更新率  [ ] 负载'),
('12. DDS 问题','【推导问题】','P1','发生概率：★★★☆☆  影响程度：★★★★☆  排查难度：★★★☆☆','AD9850/AD9833 无输出、频率不对、扫频跳变、杂散明显。','串行时序/FSYNC 错；频率字计算溢出；参考时钟假设错；复位未完成；输出滤波/偏置不足。','示波器/频率计查 DDS 原始输出；读回或打印频率字；测参考时钟。','1) 复位；2) 设固定低频；3) 检查时序与字节序；4) 校准 MCLK；5) 再做扫频。','频率字使用 64 位中间量；扫频等待稳定时间；输出端加合适低通和缓冲。','推荐：固定频率先校准，再执行扫频表。','备用：DAC DMA 生成低频保底正弦。','[ ] Reset  [ ] MCLK  [ ] 频率字  [ ] SPI/并口时序  [ ] 输出滤波'),
('13. VGA / PGA / 程控放大','【推导问题】','P1','发生概率：★★★☆☆  影响程度：★★★★☆  排查难度：★★★☆☆','AD603 增益不准、自动量程振荡、强信号削顶、弱信号噪声大。','控制 DAC 比例错；增益未标定；切档太频繁；输出/ADC 饱和；稳定时间不足。','显示控制电压、估计增益、饱和计数；示波器测输入/输出。','1) 固定档校准；2) 建立增益表；3) 加入上下阈值迟滞；4) 切档后等待稳定；5) 仅在帧边界换档。','自动量程以 20%~80% ADC 满量程为目标，并设置最小驻留帧数。','推荐：人工档位先完成，再开自动。','备用：固定中档，牺牲动态范围保稳定。','[ ] 控制电压  [ ] 增益表  [ ] 饱和  [ ] 驻留时间  [ ] 切档迟滞'),
('14. 运放相关问题','【推导问题】【比赛高频】','P1','发生概率：★★★★☆  影响程度：★★★★☆  排查难度：★★★★☆','输出贴边、波形三角化、增益低、无输出或自激。','输入共模/输出摆幅超范围；单电源却需负摆幅；GBW/压摆率不足；反馈/去耦布局差；负载太重。','示波器同时看输入、输出和电源轨；改变幅度/频率看何时失真。','1) 查供电；2) 查 DC 工作点；3) 降频降幅；4) 查闭环增益与 GBW；5) 检查去耦/反馈/负载。','给单电源交流信号加中点偏置；选够快的运放；输出串电阻隔离容性负载。','推荐：先验证 DC 点和小信号闭环增益。','备用：降低增益/带宽或改双电源。','[ ] 电源轨  [ ] 共模  [ ] 输出摆幅  [ ] SR/GBW  [ ] 去耦'),
('15. 比较器 / 波形整形','【推导问题】','P1','发生概率：★★★★☆  影响程度：★★★☆☆  排查难度：★★★☆☆','LM393 输出毛刺、边沿抖动、正弦测频不稳或高频无翻转。','无迟滞；上拉缺失/过大；门限接近噪声；开集输出误解；器件速度不足。','示波器看模拟输入和数字输出；逐步改变门限与迟滞。','1) 查输出上拉；2) 设置中点门限；3) 加迟滞；4) 查传播延迟；5) 再接 Timer。','LM393 必须确认开集上拉；高频换高速比较器或采用 ADC/FFT。','推荐：施密特迟滞 + 合理上拉。','备用：软件零交叉插值。','[ ] 上拉  [ ] 门限  [ ] 迟滞  [ ] 速度  [ ] 共地'),
('16. 模拟滤波器','【推导问题】','P2','发生概率：★★★☆☆  影响程度：★★★☆☆  排查难度：★★★★☆','截止频率不对、幅度衰减、相位/群时延异常、接负载后参数变。','R/C 误差；运放带宽不足；源/负载阻抗影响；拓扑/元件接错；没有考虑相位。','扫频画幅相曲线；用网络分析/函数源+双通道 ADC 比对。','1) 检查元件值；2) 单级验证；3) 加高阻测量；4) 扫频；5) 做频响补偿。','幅值测量前做增益校准；相位/群时延题保留整条链路的基线。','推荐：先用一阶/二阶稳定拓扑，重视相位记录。','备用：数字滤波承担可变部分。','[ ] R/C  [ ] 运放GBW  [ ] 源负载  [ ] 截止频率  [ ] 相位'),
('17. 数字滤波器','【推导问题】','P2','发生概率：★★★☆☆  影响程度：★★★☆☆  排查难度：★★★★☆','输出全零/溢出、频率响应错、延迟过大、CPU 来不及。','系数 Fs 版本错；定点缩放错；状态数组未保留；IIR 数值不稳；块长不匹配。','对单位脉冲和单频正弦做离线/串口验证；测每帧耗时。','1) 确认 Fs；2) 查系数；3) 浮点先跑通；4) 检查状态；5) 优化为 CMSIS/定点。','FIR 用于线性相位，IIR 用于低算力；滤波放在算法前要计算总延迟。','推荐：小阶 IIR/FIR，先可重复再提阶。','备用：关闭滤波，仅做多帧平均。','[ ] Fs  [ ] 系数  [ ] state  [ ] 溢出  [ ] 延迟  [ ] 耗时'),
('18. 屏幕显示','【历史真实问题】【比赛高频】','P1','发生概率：★★★★☆  影响程度：★★★☆☆  排查难度：★★★☆☆','ILI9341/OLED 黑屏、花屏、刷新慢或一刷屏 ADC 就异常。','SPI 模式/CS/DC/复位错；背光未开；共享 SPI 无锁；整屏 FillScreen 太频繁；显示函数名/API 层混用。','先画单像素/纯色；检查背光与 DC/CS；示波器看 SPI；确认调用的是头文件声明的 TFT_ILI9341_* API。','1) 背光；2) 复位/初始化；3) 固定填色；4) 单线；5) 小区域刷新；6) 最后绘制波形。','显示与采集分时，限制帧率；波形只擦局部；SPI 共享时加锁。','推荐：采集优先，UI 5~20 Hz 局部刷新。','备用：串口输出数值，关闭图形显示。','[ ] 背光  [ ] SPI模式  [ ] DC/CS  [ ] API来自头文件  [ ] 刷新率'),
('19. 按键 / 键盘','【推导问题】','P2','发生概率：★★★☆☆  影响程度：★★☆☆☆  排查难度：★★☆☆☆','按一次触发多次、长按误判、矩阵键盘串键。','未消抖；扫描太快/太慢；输入上下拉错误；阻塞式延时影响实时采样；矩阵鬼键。','串口输出原始 GPIO 与去抖事件；用 LED 显示按键状态。','1) 确认电平；2) 先做单键；3) 加 10~30 ms 去抖；4) 区分按下/释放/长按；5) 再接菜单。','按键用周期扫描状态机，不在中断里延时。','推荐：短按确认、长按返回、参数编辑限幅。','备用：串口命令输入。','[ ] 上拉/下拉  [ ] 原始电平  [ ] 去抖  [ ] 长按阈值  [ ] 不阻塞'),
('20. SPI / I2C / UART 通信','【比赛高频】','P1','发生概率：★★★★☆  影响程度：★★★★☆  排查难度：★★★☆☆','外设没响应、I2C 卡死、UART 乱码、SPI 读写全 FF/00。','CPOL/CPHA/速率错；CS 时序错；I2C 无上拉/地址错；UART 时钟/波特率错；总线被多个驱动同时用。','逻辑分析仪或示波器看 SCK/MOSI/MISO、SDA/SCL；串口回环。','1) 查供电/GND；2) 降速；3) 查模式/地址；4) 读固定 ID；5) 查片选；6) 再启用 DMA。','SPI 先 blocking 单字节；I2C 加超时和总线恢复；UART 用已知波特率回环。','推荐：拿到“一个确定读回值”再接应用。','备用：用软件 GPIO 低速验证接线。','[ ] 共地  [ ] 模式  [ ] CS/地址  [ ] 上拉  [ ] 波特率  [ ] 降速'),
('21. DMA 问题','【比赛高频】【致命】','P0','发生概率：★★★★☆  影响程度：★★★★★  排查难度：★★★★☆','DMA 不启动、只传一次、数据重复、缓冲覆盖、完成中断丢失。','请求源错；长度单位错（样本/字节）；内存对齐或地址非法；循环/乒乓配置错；CPU 与 DMA 同时改缓冲。','在 DMA 完成 ISR 翻转 GPIO；打印块序号和首尾样本；检查状态寄存器/错误标志。','1) 单次短传输；2) 确认请求源；3) 检查长度/宽度；4) 完成中断；5) 再循环/双缓冲。','生产者 DMA、消费者算法必须拥有不同块；用块序号确认一致快照。','推荐：先 one-shot，再 Ping-Pong。','备用：降低采样率/块长，或暂用轮询 ADC。','[ ] 请求源  [ ] 长度单位  [ ] 地址/对齐  [ ] 完成标志  [ ] 缓冲所有权'),
('22. Timer 问题','【比赛高频】','P1','发生概率：★★★★☆  影响程度：★★★★☆  排查难度：★★★☆☆','PWM/触发频率不对、输入捕获跳变、ADC 实际 Fs 偏差。','时钟源假设错；prescaler/ARR 装载错；更新事件未产生；中断/事件路由错；计数器溢出。','在 GPIO 输出 timer PWM/trigger；读 timer clock、PSC、ARR；用示波器实测。','1) 确认 timer clock；2) 计算并打印目标/实际；3) 实测输出；4) 检查 event；5) 最后连接 ADC/DMA。','把 Fs=timer_clk/((PSC+1)(ARR+1)) 写进注释/串口；改变一个参数后重新计算。','推荐：独立 timer 作为全链路唯一采样时基。','备用：固定预置频率并查表。','[ ] timer时钟  [ ] PSC  [ ] ARR  [ ] update事件  [ ] 实测Fs'),
('23. 中断问题','【比赛高频】','P1','发生概率：★★★★☆  影响程度：★★★★☆  排查难度：★★★★☆','系统卡死、丢采样、变量偶发异常、UI 卡顿。','ISR 太长；优先级反转；共享变量无 volatile/临界区；高频中断超过 CPU；在 ISR 内 SPI/printf。','GPIO 标记 ISR 宽度；统计中断计数；关掉单个 IRQ 二分定位。','1) 看 ISR 时间；2) 只置标志；3) 主循环处理；4) 检查 volatile/原子性；5) 设置优先级。','ISR 仅搬运/置标志；DMA 完成优先于 UI；printf 移到低优先级任务。','推荐：采样 ISR/DMA 最高，算法次之，UI 最低。','备用：降低刷新/采样率。','[ ] ISR时长  [ ] volatile  [ ] 优先级  [ ] printf不在ISR  [ ] 计数器'),
('24. RAM / Flash / CPU 性能','【比赛高频】','P0','发生概率：★★★★☆  影响程度：★★★★★  排查难度：★★★★☆','FFT 后崩溃、栈溢出、DMA 数组破坏、帧处理跟不上。','大数组放栈；双通道+FFT 缓冲估算不足；浮点/绘图耗时高；堆碎片；链接段超限。','查看 .map；打印数组地址/剩余 RAM；测一帧采集、算法、显示耗时。','1) 查 map；2) 把大数组移全局静态；3) 减少 N/双缓冲；4) 降 UI；5) 再优化 DSP。','提前列 RAM 预算：ADC 原始+工作+FFT+显示；DMA 缓冲静态对齐。','推荐：保留 RAM/CPU 30% 余量。','备用：缩小 FFT N、只做单通道/低刷新。','[ ] map  [ ] 大数组非栈  [ ] DMA内存  [ ] 帧耗时  [ ] 余量'),
('25. FPGA + MCU 协同','【推导问题】','P2','发生概率：★★★☆☆  影响程度：★★★★★  排查难度：★★★★★','FIFO 溢出、数据错位、偶发丢块、时钟域后随机错误。','跨时钟域未同步；FIFO 水位策略错；帧同步缺失；MCU 总线吞吐不足；复位顺序错。','输出帧头/块号/CRC；观察 FIFO empty/full；逻辑分析仪看握手。','1) FPGA 先产生已知递增数据；2) 验证 FIFO；3) MCU 读块校验；4) 压力测试；5) 再接 ADC。','所有块带序号和长度；用 FIFO 水位/背压；复位后等待时钟稳定。','推荐：FPGA 采样缓存，MCU 只取完整块。','备用：降低 FPGA 输出率或增大块缓存。','[ ] 时钟域  [ ] FIFO水位  [ ] 帧头/序号  [ ] 吞吐  [ ] 复位顺序'),
('26. 模块拼装与资源冲突','【历史真实问题】【比赛高频】【致命】','P0','发生概率：★★★★★  影响程度：★★★★★  排查难度：★★★★☆','单模块都好用，组合后某一模块失效/编译宏冲突/无中断。','Timer/DMA/SPI/GPIO/IRQ 重复分配；SysConfig 实例名不一致；两个模块修改同一生成资源。','建立资源表，逐项查 .syscfg 和 README；用最小组合二分排查。','1) 固定主时基；2) 分配 ADC/DMA；3) 分配 SPI；4) 分配 GPIO/IRQ；5) Generate/Build；6) 再接功能。','一项硬件资源只能有一个 owner；功能模块由适配层共享而非各自重配外设。','推荐：赛前冻结资源矩阵。','备用：关掉非关键模块，腾出 timer/DMA/SPI。','[ ] Timer  [ ] DMA  [ ] SPI  [ ] GPIO  [ ] IRQ  [ ] SysConfig实例名'),
('27. 模块库调用 / API 迁移','【历史真实问题】【比赛高频】','P0','发生概率：★★★★★  影响程度：★★★★☆  排查难度：★★★☆☆','复制模块后 API 未声明、函数名看似接近但无法链接，或示例变量耦合难拆。','只看 .c 未看 .h；把平台适配 API 与通用驱动 API 混用；漏复制 .inc/依赖；新 .c 未加入 Build。','以头文件为准搜索声明；对照 README 的“复制清单→include→init→process→get_result”。','1) 复制 README 指定文件；2) Refresh；3) include 主 .h；4) Build；5) 只调用已声明 API；6) 再接 main。','禁止自行补写未声明 extern；用统一 status 返回码处理失败；保留模块版本/来源。','推荐：每个模块做 20 行最小 smoke test。','备用：回到 DriverLib 直接实现最小功能。','[ ] 头文件声明  [ ] 依赖文件  [ ] 已参与Build  [ ] init返回码  [ ] 无虚构API'),
('28. 完整数据链','【历史讨论问题】【比赛高频】','P0','发生概率：★★★★☆  影响程度：★★★★★  排查难度：★★★★☆','“结果错”但不知道错在 ADC、预处理、算法还是显示。','跳过原始数据验证；算法输入与采集缓冲不同步；标定/单位散落；UI 显示旧帧。','每层输出可观测量：ADC min/max、块序号、去 DC 后 RMS、FFT peak、最终值/时间戳。','1) 信号源；2) 模拟输入；3) ADC 原始数组；4) DMA 块；5) 预处理；6) 算法；7) UI。','给每帧带 sequence；所有结果关联同一 sequence；保留串口诊断开关。','推荐：ADC→预处理→算法→结果四层接口。','备用：串口只输出 ADC/算法中间量，不画图。','[ ] 同一帧号  [ ] 原始样本  [ ] 单位  [ ] 标定  [ ] 中间变量  [ ] UI时间戳'),
('29. 电源 / 接地 / 模拟硬件','【推导问题】【比赛高频】【致命】','P0','发生概率：★★★★★  影响程度：★★★★★  排查难度：★★★☆☆','噪声陡增、接屏闪烁、参考漂移、系统重启、测幅/测相不稳。','电流余量不足；地回流穿过模拟前端；ADC/DAC reference 去耦差；DC-DC 开关纹波；屏幕瞬态电流。','万用表查 DC，示波器短地弹簧查纹波；断开大负载对比噪声谱。','1) 电源余量；2) 共地；3) 每轨去耦；4) 检查 reference；5) 分离大电流回路；6) 再调算法。','数字/模拟地在单点低阻连接；敏感参考旁加专用去耦；屏幕电源加储能。','推荐：先把“静态 DC + 噪声”量出来再上算法。','备用：外接稳压电源、关闭背光/高功耗输出。','[ ] 电流余量  [ ] 共地  [ ] 去耦  [ ] reference  [ ] DC-DC纹波'),
('30. 信号完整性','【推导问题】','P1','发生概率：★★★☆☆  影响程度：★★★★☆  排查难度：★★★☆☆','高频幅度变小、振铃、不同接法结果不同、示波器一接就变。','50 Ω/高阻假设不一致；飞线/面包板寄生；探头 1x 负载大；地夹太长；模块串联加载。','换 10x 探头、短地弹簧、短线；在每一级测幅；检查源/负载阻抗。','1) 测源端；2) 测前端输入；3) 测 ADC 前；4) 比较接入前后；5) 必要时端接。','高频用同轴/短线；明确每级输入输出阻抗；避免长地夹。','推荐：先把模拟链逐级测通。','备用：降频并采用高阻缓冲。','[ ] 探头倍率  [ ] 地夹  [ ] 50Ω  [ ] 线长  [ ] 逐级幅度'),
('31. 精度不够时怎么办','【推导问题】【比赛高频】','P1','发生概率：★★★★☆  影响程度：★★★★☆  排查难度：★★★★☆','测频、测幅、测相、THD 或群时延离题目指标差一点。','把系统误差当算法误差；采样不足；未校准；直接套高阶插值但条件不满足。','先用已知标准源建立误差随频率/幅值变化的曲线，再决定优化层级。','1) ADC/时钟/通道基线；2) 增加观察时间；3) 去噪/平均；4) 插值/拟合；5) 频响补偿；6) 复测。','测频：三点抛物线→Jacobsen/Quinn→CZT；测幅：校准+RMS/拟合；测相：延时校准+unwrap；THD：相干采样+窗校正。','推荐：先标定、再多周期平均、最后加插值。','备用：降低指标展示范围，切换可靠的低阶算法。','[ ] 标准源  [ ] 时钟  [ ] 校准  [ ] 多周期  [ ] 采样N  [ ] 算法适用条件'),
('32. 比赛现场应急降级','【比赛高频】【快速修复】','P0','发生概率：★★★★☆  影响程度：★★★★★  排查难度：★★☆☆☆','复杂方案调不通、时间不足、系统勉强运行但不稳定。','过早追求全功能；同时改硬件、SysConfig、算法；没有保底路径。','每个模块记录“最优/次优/保底”；比赛中只做一个变量的改动并保存可运行版本。','1) 冻结能跑版本；2) 确定核心评分项；3) 关 UI/高级算法；4) 固定量程/参数；5) 再逐项恢复。','FFT 不稳→Timer；高级插值不稳→三点抛物线；自动量程失败→固定档；Ping-Pong 异常→单帧 DMA；UI 卡→串口。','推荐：先交付稳定核心结果。','备用：手工配置+最小采样链+简化显示。','[ ] 可运行备份  [ ] 核心指标  [ ] 关闭非关键功能  [ ] 版本号  [ ] 复测')]

SYMPTOMS = [
('不上电','限流、短路、反接、供电余量','1、29'),('电源反复重启','浪涌、屏幕/DDS 负载、DC-DC','1、29'),('烧录器找不到芯片','接线、目标电压、复位','2'),('下载成功但不运行','启动/HardFault/时钟','2、3'),('SysConfig 宏缺失','未生成、实例名错','3'),('编译找不到头文件','include 路径/复制清单','3、27'),('链接 undefined reference','漏 .c 或 API 名不匹配','3、27'),('ADC 全 0','输入悬空、触发/DMA','4'),('ADC 满量程','输入范围/偏置/饱和','4'),('ADC 波形乱','共地、采样时钟、DMA','4、29'),('双通道不同步','非同一触发、块序号不同','4、8、21'),('频率高时不准','Timer 分辨率/比较器速度','5、15、22'),('低频刷新太慢','周期数太多、窗口太长','5、31'),('正弦测频跳','门限抖动、DC、无迟滞','5、15'),('FFT 主峰错','Fs/N、搜索范围、混叠','6'),('FFT 频率跳','泄漏、栅栏效应、信噪低','6、31'),('FFT 幅值偏低','窗增益、标定、削顶','6、7'),('THD 很大','泄漏、削顶、谐波/噪声','6、31'),('Vpp 跳动','尖峰、噪声、分位数缺失','7'),('RMS 不对','DC 是否去除、AC/DC 定义','7'),('相位跳 180°','wrap/unwrap、相位符号','8'),('相位随频率漂移','通道延迟、模拟链相移','8、16'),('李萨如图不闭合','XY 不同步、频率比不稳、缩放','4、8、18'),('波形识别误判','噪声、阈值、饱和','9'),('猝发没抓全','无预触发、环形缓冲覆盖','10、21'),('DAC 无输出','未使能、码值/Vref、触发','11'),('DDS 频率错误','频率字、MCLK、时序','12'),('DDS 杂散大','时钟、滤波、布局','12、30'),('增益自动来回跳','无迟滞、稳定时间不足','13'),('运放输出贴边','共模/摆幅/偏置','14'),('运放高频三角化','压摆率、GBW','14'),('比较器毛刺','无迟滞、上拉/门限','15'),('滤波后相位不对','群时延、负载、拓扑','16'),('数字滤波输出全零','系数、缩放、状态','17'),('屏幕黑','背光、SPI、初始化','18'),('屏幕花','SPI 模式/DC/CS、共享总线','18、20'),('一刷屏 ADC 就乱','CPU/SPI 抢占、供电','18、23、29'),('按键多触发','去抖、扫描','19'),('SPI 无响应','CPOL/CPHA、CS、供电','20'),('I2C 卡死','上拉、地址、总线恢复','20'),('UART 乱码','波特率/时钟/地','20'),('DMA 只传一次','循环配置、完成标志','21'),('DMA 数据覆盖','双缓冲所有权、长度','21'),('PWM/Fs 不对','timer clock、PSC/ARR','22'),('中断丢失','优先级、ISR 太长','23'),('系统偶发卡死','ISR/共享变量/栈','23、24'),('FFT 后崩溃','RAM/栈不足','24'),('FPGA 数据错位','FIFO/时钟域','25'),('单模块好组合坏','资源冲突','26'),('API 名字相近但编不过','未看 .h、适配层混用','27'),('结果显示旧数据','帧号/数据链不同步','28'),('接线后幅度变了','阻抗/探头/长线','30'),('精度差一点','未标定、窗口、插值','31'),('时间不够调不通','未降级、改动太多','32')]

NEEDS = [
('测频','Timer 输入捕获/计数 → 频率','Timer、比较器','周期法/计数法','FFT+插值'),('正弦高精度测频','ADC → 去DC/窗 → FFT','ADC DMA、CMSIS FFT','三点抛物线','CZT/Quinn'),('测 Vpp','ADC → 标定 → 分位数','ADC、校准','robust Vpp','min/max'),('测 Vrms','ADC → 去DC → RMS','ADC、CMSIS','AC RMS','正弦拟合'),('测相位','双 ADC 同触发 → FFT 复相位','Dual ADC、FFT','相位差+校准','零交叉'),('李萨如图','双 ADC 同帧 → 缩放 → TFT','Dual ADC、TFT','逐点连线','抽样点显示'),('识别波形','ADC → 特征/谐波 → 分类','ADC、FFT','时域+谐波规则','人工判读'),('检猝发','ADC 环形缓冲 → 触发 → 快照','DMA、比较器','预触发+后触发','软件门限'),('扫频','DDS → DUT → ADC','DDS、ADC','频点表+稳定等待','手动频点'),('测幅频','DDS → DUT → ADC 幅值','DDS、ADC、校准','RMS/拟合','Vpp'),('测相频','DDS → DUT → 双 ADC','DDS、Dual ADC','FFT 相位','零交叉'),('测群时延','扫频相位 → unwrap → 差分','DDS、Dual ADC','相位导数','有限差分平滑'),('测 THD','ADC → FFT → 谐波比','ADC、FFT','相干采样+窗','锁定基波 bin'),('测 GBW','扫频 → 增益曲线','DDS、ADC','-3 dB 搜索','手工频点'),('测压摆率','阶跃 → ADC → 斜率','DAC/DDS、ADC','线性段拟合','两点差分')]

def setup(doc):
    sec=doc.sections[0]; sec.top_margin=Inches(1); sec.bottom_margin=Inches(1); sec.left_margin=Inches(1); sec.right_margin=Inches(1); sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Microsoft YaHei'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
    for name,size,color,before,after in [('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',11.5,DARK,10,5)]:
        s=styles[name]; s.font.name='Microsoft YaHei'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; fonts(header.add_run('电子设计竞赛单片机 / 信号题现场手册'),8.5,'666666')
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; fonts(footer.add_run('内部比赛速查版本｜基于项目历史对话与工程资料整理'),8,'666666')

def title(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(90); p.paragraph_format.space_after=Pt(14); fonts(p.add_run('电子设计竞赛单片机 / 信号题'),25,DARK,True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(28); fonts(p.add_run('比赛现场问题与解决方法大全'),21,BLUE,True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; fonts(p.add_run('面向 MSPM0 / STM32 / ADC / FFT / DDS / 显示与系统集成'),11,'666666')
    add_callout(doc,'使用原则','先看现象反查表；再从 L0 电源向上排查；最后才修改算法。每一次只改一个变量，并保留最近一次可运行版本。')
    doc.add_page_break()

def intro(doc):
    heading(doc,'使用范围与历史检索结论',1)
    add_text(doc,'本手册完成了当前项目可访问历史聊天的全量扫描：22 份 JSONL 会话、658 条记录，其中包含 24 条用户记录与 128 条助手记录。历史会话集中在 MSPM0G3507 双路同步 ADC、ILI9341 TFT 与李萨如图集成，时间为 2026-08-16 至 2026-08-17。')
    add_callout(doc,'真实性边界','【历史真实问题】仅指历史聊天/日志确实出现的工程或工具问题，并不等价于硬件上板实测。历史中未找到“HDN3-5D05A1、±250 mA、0.6~1 A 跳动、串口屏闪烁”等原始对话证据；这些仅按【推导问题】收录为电源高风险项。')
    add_bullets(doc,['【历史真实问题】双路同步 ADC + ILI9341 李萨如集成；要求仅改 main.c；SysConfig 被声明已配好。','【历史真实问题】API 层混用：历史记录中出现过未在活动头文件声明的显示函数；正确做法是以 .h 声明为准，调用 TFT_ILI9341_* 与 SignalTFTILI9341_MSPM0_Init。','【历史真实问题】工程操作失败：精确替换 old_string 不匹配（5 次）、未读全文件便覆盖被拒绝（3 次）、文件不存在（5 次）；这直接转化为比赛模块移植规则。','【历史真实问题】运行环境超时/断连：StreamInactivityTimeoutError 14 次、APIConnectionError 4 次、ECONNRESET 3 次；这是工具链风险，不作为板端故障。'])
    heading(doc,'标签与优先级',2); add_text(doc,'P0：赛前必须解决；P1：必须知道排查；P2：有时间优化；P3：低概率。标签中的【比赛高频】表示比赛结构下高概率，并非历史统计。')
    heading(doc,'分层调试模型 L0-L8',2)
    t=doc.add_table(rows=1, cols=3); set_table_widths(t,[1100,2100,6160]);
    for c,x in zip(t.rows[0].cells,['层','对象','本层最小证据']): shade(c,PALE); set_cell(c,x,True,DARK)
    layers=[('L0','电源/地','各电源轨正确、限流未触发'),('L1','模拟输入','ADC 脚处波形和偏置正确'),('L2','ADC','固定 DC 与低频正弦码值正确'),('L3','搬运','DMA 块序号连续、无覆盖'),('L4','原始数据','min/max/mean 与示波器一致'),('L5','预处理','去 DC、标定、窗/滤波可观察'),('L6','算法','峰值、相位、特征值正确'),('L7','后处理','单位、校准、平均正确'),('L8','UI/通信','显示的是同一帧的新结果')]
    for a,b,c in layers:
        r=t.add_row().cells
        for cell,x in zip(r,[a,b,c]): set_cell(cell,x)
    heading(doc,'系统出了问题后的 10 分钟流程',2)
    add_bullets(doc,['0~1 分钟：断开扩展件，检查电源、电流、共地。','1~2 分钟：最小 MCU LED/串口验证。','2~4 分钟：示波器确认模拟输入与 ADC 脚。','4~5 分钟：打印 ADC 原始 min/max/首尾样本。','5~6 分钟：确认 Timer Fs、DMA 完成与帧号。','6~8 分钟：检查预处理与算法中间量。','8~9 分钟：对比标准信号源，检查校准。','9~10 分钟：最后才恢复 UI、扫频或高级算法。'])
    add_callout(doc,'禁止项','不要一看到最终结果错误就先改 FFT、滤波或分类阈值；先证明 L0~L5 正常。')

def cards(doc):
    heading(doc,'按系统模块的故障排查卡',1)
    for title_,tags,pri,rating,phen,causes,quick,steps,fix,reco,backup,check in CARDS:
        heading(doc,title_,2)
        add_text(doc,f'{tags}  比赛优先级：{pri}  {rating}')
        add_text(doc,'现象：'+phen,'现象：')
        add_text(doc,'高概率原因：'+causes,'高概率原因：')
        add_text(doc,'快速判断：'+quick,'快速判断：')
        add_text(doc,'排查顺序：'+steps,'排查顺序：')
        add_text(doc,'解决方法：'+fix,'解决方法：')
        add_text(doc,'推荐方案：'+reco,'推荐方案：')
        add_text(doc,'备用方案：'+backup,'备用方案：')
        add_text(doc,'现场快速检查：'+check,'现场快速检查：')

def indexes(doc):
    heading(doc,'我看到这个现象，应该查哪里？',1)
    add_text(doc,'症状反查表：先按“最可能原因”快速进入章节；若仍未定位，回到 L0~L8 分层模型。')
    t=doc.add_table(rows=1,cols=3); set_table_widths(t,[2600,4700,2060])
    for c,x in zip(t.rows[0].cells,['现象','先查/最可能原因','对应章节']): shade(c,PALE); set_cell(c,x,True,DARK)
    for a,b,c in SYMPTOMS:
        r=t.add_row().cells
        for cell,x in zip(r,[a,b,c]): set_cell(cell,x)
    heading(doc,'比赛需求 → 推荐实现方案',1)
    t=doc.add_table(rows=1,cols=5); set_table_widths(t,[1650,2450,1800,1800,1660])
    for c,x in zip(t.rows[0].cells,['需求','推荐数据链','所需模块','推荐算法','备用算法']): shade(c,PALE); set_cell(c,x,True,DARK,8.5)
    for row in NEEDS:
        r=t.add_row().cells
        for cell,x in zip(r,row): set_cell(cell,x,size=8.5)
    heading(doc,'精度增强决策树',1)
    add_bullets(doc,['测频不准 → 先查 Fs/Timer 时钟 → 增加观测周期 → 三点抛物线 → Jacobsen/Quinn → CZT/Zoom FFT。','测幅不准 → 先 ADC offset/gain 标定 → 检查削顶 → 去 DC → 分位数/MAD → 正弦拟合。','测相不准 → 双路同触发 → 同源零相位基线 → 固定延时校准 → FFT 复相位 → unwrap/频响补偿。','THD 不准 → 先防混叠和削顶 → 尽量相干采样 → 窗与 coherent gain 修正 → 谐波 bin 一致化。','群时延不准 → 相位先 unwrap → 频率点加密 → 平滑差分 → 扣除系统基线。'])
    heading(doc,'最优 → 次优 → 保底的降级表',1)
    t=doc.add_table(rows=1,cols=4); set_table_widths(t,[2250,2370,2370,2370])
    for c,x in zip(t.rows[0].cells,['故障场景','最优方案','次优方案','保底方案']): shade(c,PALE); set_cell(c,x,True,DARK)
    rows=[('FFT 频率不稳定','加窗+插值+校准','三点抛物线','Timer/零交叉'),('DMA 连续采集异常','Ping-Pong+帧号','单次 DMA','轮询短帧'),('自动量程振荡','迟滞+驻留','固定几档','固定中档'),('屏幕影响实时性','局部低帧率刷新','仅数值刷新','串口输出'),('双通道相位不稳','同步 ADC+校准','零交叉','只报幅频'),('复杂分类不稳','时频融合','时域规则','人工识别/未知'),('DDS 扫频异常','频点表+稳定检测','低速扫频','手动固定频点')]
    for row in rows:
        r=t.add_row().cells
        for cell,x in zip(r,row): set_cell(cell,x)

def quick(doc):
    doc.add_page_break(); heading(doc,'比赛现场故障速查表（30 秒定位）',1)
    add_callout(doc,'使用法','看到现象，先做“先检查”，再用“最快解决”保住核心功能；详细原因回查前文对应章节。')
    quickrows=[
('不上电/重启','断外设量 3.3/5 V 与电流','短路、浪涌、供电不足','最小系统→逐个回接'),('MCU 不运行','只接调试器跑 LED','SWD/Reset/启动异常','Mass erase+最小程序'),('烧录失败','目标电压与调试线','接反、外供冲突','断外设、降速下载'),('ADC 无数据','DC 输入+打印原始码','触发/DMA/输入悬空','单次短 DMA'),('ADC 异常','min/max、ADC脚示波器','饱和、共地、Fs','加偏置/限幅/重查 Timer'),('DMA 异常','块号、ISR GPIO','长度/请求源/覆盖','one-shot DMA'),('FFT 异常','Fs/N、去DC、最大bin','泄漏/混叠/窗','Hann+三点插值'),('测频异常','Timer 时钟与边沿','门限抖动/周期太少','Timer→FFT 复核'),('测幅异常','校准与削顶计数','offset/gain/尖峰','分位数 Vpp+RMS'),('测相异常','同源零相位基线','不同步/延时','延时校准+unwrap'),('DDS 无输出','Reset、MCLK、频率字','时序/字节序','固定低频验证'),('DAC 无输出','中码、电压、触发','未使能/Vref','固定码→低频'),('SPI 不通','模式、CS、降速','CPOL/CPHA/DC','读固定 ID'),('I2C 不通','上拉、地址、SCL','无上拉/总线锁','总线恢复+低速'),('UART 不通','回环、波特率、GND','时钟/电平','固定波特率回环'),('屏幕不亮','背光、纯色、SPI','初始化/DC/CS','关闭 UI 改串口'),('按键没反应','GPIO 原始电平','上下拉/扫描','单键状态机'),('系统卡死','ISR GPIO、关闭外设','ISR过长/栈','ISR只置标志'),('CPU 来不及','测每帧耗时','FFT/绘图过重','降 N/刷新率'),('RAM 不足','.map 与数组地址','栈大数组/双缓冲','数组移全局、减 N'),('模块冲突','资源矩阵','Timer/DMA/SPI 重复','关闭非关键模块'),('高频波形失真','探头/线/运放输出','SR/GBW/阻抗','降频/短线/缓冲')]
    t=doc.add_table(rows=1,cols=4); set_table_widths(t,[2100,2450,2450,2360])
    for c,x in zip(t.rows[0].cells,['现象','先检查','最可能原因','最快解决']): shade(c,PALE); set_cell(c,x,True,DARK)
    for row in quickrows:
        r=t.add_row().cells
        for cell,x in zip(r,row): set_cell(cell,x)
    heading(doc,'赛前 15 分钟验收清单',1)
    add_bullets(doc,['最小 MCU（LED/串口）可运行且可重新烧录。','各电源轨、电流、正负电源和共地记录完成。','ADC 固定 DC、低频正弦、双通道同步均已验证。','Fs 用示波器实测，DMA 块号连续且无覆盖。','标准信号源下测频/测幅/测相的校准基线已保存。','TFT 可画纯色/线段，但 UI 降级为可关闭。','DDS/DAC 固定频率/固定码可验证；扫频有稳定等待。','资源矩阵、可运行工程备份和降级路径已准备。'])
    heading(doc,'历史方案冲突与裁决',1)
    add_text(doc,'历史聊天中曾出现“自行声明平台级显示/ADC 函数”与“仅调用实际头文件 API”两种做法。前者导致 API 未声明/虚构风险，应淘汰；当前比赛工程应以 modules/*.h 为唯一契约：SignalDualADC_* 负责采集，SignalTFTILI9341_MSPM0_Init 绑定平台，TFT_ILI9341_* 负责绘制。')
    heading(doc,'附：本次历史检索覆盖说明',1)
    add_text(doc,'已检索项目内所有匹配 chats/ 的 JSONL 历史会话。历史会话以工程集成与工具验证为主，不含对 AD7606、AD9226、AD603、HDN3-5D05A1 或实际电源电流波形的可确认对话；相关章节因此使用【推导问题】。项目 README、模块 README、双 ADC 与 ILI9341 头文件用于核实当前工程的真实模块边界。')

doc=Document(); setup(doc); title(doc); intro(doc); cards(doc); indexes(doc); quick(doc); doc.save(OUT)
print(OUT.resolve())
