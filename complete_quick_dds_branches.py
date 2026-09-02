from docx import Document
from docx.shared import Pt

PATH = '电子设计竞赛单片机_信号题_比赛现场速查手册.docx'
doc = Document(PATH)
for table in doc.tables:
    if len(table.rows) < 2:
        continue
    probe = table.rows[1].cells[0].text
    if '测 Timer/DMA 更新率、DAC 引脚' in probe:
        cells = table.add_row().cells
        cells[0].text = '正常 / 异常'
        cells[1].text = '正常：Fupdate 与 DAC/滤波后频率一致。异常：比例错查 Fupdate；无输出查 Timer/DMA/DAC。'
        cells[2].text = '异常即回到固定 1 kHz，不进入扫频。'
    if 'AD9833：MCLK/FSYNC/SCLK/SDATA/RESET' in probe:
        cells = table.add_row().cells
        cells[0].text = '正常 / 异常'
        cells[1].text = '正常：参考时钟、装载时序、输出逐级成立。异常：比例错查 MCLK/CLKIN；无输出查 RESET/接口。'
        cells[2].text = '先固定频点；AD9833/AD9850 不混用。'
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7.6)
doc.save(PATH)
print('updated', PATH)
