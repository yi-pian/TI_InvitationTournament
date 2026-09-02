from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT_FULL = '电子设计竞赛单片机_信号题_完整Debug手册.docx'
OUT_QUICK = '电子设计竞赛单片机_信号题_比赛现场速查手册.docx'

NAVY = '17365D'; BLUE = '1F4E78'; TEAL = '0F6B78'; RED = 'B00020'; GRAY = 'F2F5F7'; LIGHT = 'D9EAF7'

RESOURCE_ROWS = [
    ('CPU', '32 MHz', '当前工程 CPUCLK_FREQ=32000000'),
    ('Timer', 'TIMG0', '双 ADC 的定时触发源；500 kSPS 时计数约 64'),
    ('ADC0', 'PA25 / MEM0 ch.2', '通道 A，VDDA≈3.3 V，DMA CH0'),
    ('ADC1', 'PA17 / MEM0 ch.2', '通道 B，VDDA≈3.3 V，DMA CH1'),
    ('DMA', 'CH0 / CH1', '分别搬运 ADC0 / ADC1；模块拥有 DMA_IRQHandler'),
    ('SPI', 'SPI1', 'ILI9341 总线'),
    ('TFT GPIO', 'PB9/PB8/PB6/PB15/PB12', 'SCK/MOSI/CS/DC/背光'),
    ('SWD', 'PA20 / PA19', 'SWCLK / SWDIO；勿复用'),
]

# 此表只收录当前工程可核实对象；其它卡均使用“未实现”措辞，不补造 API。
CARDS = [
('01','电源、接地与重启','P0','当前工程相关（板级供电）','系统一接屏、扩展板或信号源就复位；串口/屏幕随机乱码。','电压掉到芯片最低工作范围、复位脚被拉低或电流异常时，STOP：先断扩展模块，不能查 SPI/算法。','最小系统只保留开发板、SWD 与 3.3 V；逐个接回 TFT、信号源、外设。','断扩展模块后测 3.3 V、5 V 和电流；用示波器看复位瞬间的压降。','3.3 V 稳定，复位不随负载变化；接入模块后仍可稳定下载与运行。','掉压：供电余量/接地回流；仍 3.3 V 但重启：复位、SWD 或时钟；只接屏异常：背光/接口短路。','检查上电初始化是否过早拉高背光，确认没有把 PA19/PA20 配成普通 GPIO。','检查地线、极性、焊桥、TFT 供电与背光电流；模拟输入不得超过 VDDA。','先减载、共地、加去耦；用外部稳压源或关闭背光验证。','无专用 API；参考开发板原理图和示波器电源启动波形。','连续运行 5 分钟、反复刷新屏幕和下载均不重启。','关闭背光、只保留数值显示与单通道采集。'),
('02','MCU 无法启动、烧录或调试','P0','当前工程真实实现','下载失败、无法 halt、程序像没运行；接外设后才出现。','SWD 不可用时 STOP：不要重配 SysConfig；先恢复最小硬件和 SWD。','仅开发板+SWD，移除外设；确认目标芯片供电。','第一测量：PA20(SWCLK)、PA19(SWDIO)、NRST 和 3.3 V。','调试器可识别芯片、可擦除/下载；PA20/PA19 不被外设占用。','无识别：供电/SWD；可下载不跑：启动文件/时钟；接外设才失败：引脚冲突或掉压。','检查 target 配置、芯片型号、Debug/Release 链接脚本；不要手改生成的 SysConfig 文件。','检查 PA20/PA19 未被复用，启动后不要重配 SWD 复用。','降低 SWD 时钟、擦除后下载；逐个恢复外设。','当前工程入口为 main.c；SWD 引脚 PA20/PA19。','连续 10 次下载均成功，断电重上电后应用能启动。','保留只点灯/只串口的最小镜像。'),
('03','CCS、SysConfig 与工程配置','P0','当前工程真实实现','编译突然大量报错、生成文件不匹配、下载后外设不响应。','SysConfig 生成文件与 .syscfg 不一致时 STOP：不要在 Debug 生成文件里“补改”。','复制工程后只保留一个 .syscfg 和一个可复现 build 配置。','第一测量：Clean 后重新生成，检查编译日志中实际包含的源文件和 include 路径。','工程能 clean build；生成的 ti_msp_dl_config.h 与当前 .syscfg 同步。','只报头文件错：路径/模块未加入；引脚错：SysConfig 未生成；能编译不工作：配置与 C 代码对象不一致。','检查 CPUCLK_FREQ=32000000、外设实例和中断符号是否来自同一套生成文件。','检查 current 工程只调用头文件中实际声明的接口。','回到 .syscfg 修改、重新生成、完整 clean build；不要手改生成代码。','当前工程：signal_contest_template.syscfg 与 Debug/ti_msp_dl_config.h。','重新打开工程、Clean/Rebuild 后无手工补丁仍可下载运行。','冻结已验证的 .syscfg，新增功能在副本工程验证。'),
('04','资源冲突、API 与模块迁移','P0','当前工程真实实现','单独模块正常，组合后编译失败、无波形、屏幕黑或中断不进。','同一外设/IRQ/DMA 通道被两处拥有时 STOP：先确定 owner，禁止两边同时初始化。','只保留一个模块和它的 README 最小例程，再合入第二个模块。','第一测量：列出 Timer、ADC、DMA、SPI、GPIO、IRQ 的 owner。','每个资源只有一个初始化者和一个 ISR owner；DMA CH0/CH1 属双 ADC 模块。','链接重定义：重复 ISR；初始化后失效：外设被二次配置；编译错：API 来自别的版本。','核对模块 .h、README 与当前工程模块清单；无声明的函数一律视为未实现。','检查 DMA_IRQHandler 不在 main.c 再定义；避免复用 PA19/PA20。','先合资源表，再删重复初始化；采用当前工程公开 API。','真实 API：SignalDualADC_Init/Start/IsFinished；TFT 初始化：SignalTFTILI9341_MSPM0_Init。','组合工程可 clean build；TFT 刷新和双 ADC 同时运行。','先保留采集+数值显示，暂停波形绘制与高级算法。'),
('05','Timer 与实际采样率 Fs','P0','当前工程真实实现','频率、FFT 横轴、相位全部成比例错误；设定 500 kSPS 但测量不对。','实际 Fs 未验证时 STOP：不要相信 FFT、频率或相位结果。','定时器触发 ADC，输入固定频率正弦或方波。','第一测量：在 Timer ISR/触发点翻转 GPIO，用示波器测周期；再由相邻采样点或捕获验证。','当前 main.c 目标 500 kSPS，CPU 32 MHz，理论 Count≈64；必须以实测为准。','Fs 偏低：时钟/分频/Count；偶发跳变：中断或重配；设定与实测不同：单位/整数取整。','核对 SignalDualADC_SetSampleRate() 的返回与 SignalDualADC_GetConfiguredRate()；定时器最大计数为 65536。','检查 Timer 时钟源、分频、触发输出和 ADC 触发选择。','以实测 Fs 更新算法参数；先固定 Fs，再恢复显示/FFT。','参考 timer_capture_minimum 与当前双 ADC 模块的 SetSampleRate/GetConfiguredRate。','示波器与软件打印 Fs 相符，连续帧频率稳定。','降低采样率、用零交叉/计数器测频。'),
('06','ADC 原始采样：全 0、满量程、乱跳','P0','当前工程真实实现','g_raw_a/g_raw_b 全 0、全 4095、固定值或无规律跳变。','raw data 不对时 STOP：不要查 DMA 后算法、FFT 或 UI。','ADC 单通道、直流 0.5/1.65/2.5 V，关闭复杂前端。','第一测量：Watch g_raw_a[0/512/1023]、g_raw_b[0/512/1023]，并在 ADC 引脚实测电压。','12 bit、Full Scale=4095、VDDA≈3.3 V：0.5 V≈620，1.65 V≈2048，2.5 V≈3102。','全 0：引脚/触发/参考；全 4095：过压或通道错；跳变：悬空、地噪声、采样时间不足。','检查 ADC0=PA25、ADC1=PA17、MEM0 channel 2、VDDA 参考；输入范围不得超过 VDDA。','检查 SignalDualADC_Init 后才 Start；用真实 g_raw_a/g_raw_b，不把建议变量当现成变量。','接三档已知电压逐项验证；必要时延长采样时间、降低源阻抗。','当前工程：SignalDualADC_Start(g_raw_a,g_raw_b,SIGNAL_SAMPLE_COUNT)。','两路三档码值均在合理误差内，连续采样无异常满量程。','固定 DC 测量或单通道采集，停止 FFT/相位。'),
('07','DMA 与 Buffer 所有权','P0','当前工程真实实现','DMA 只跑一次、帧混杂、FFT 偶发错、数组似乎没有变化。','DMA 未完成或 buffer 正在写时 STOP：不要运行 FFT/测量。','单帧 DMA 采集，先不画图、不做 FFT。','第一测量：看 SignalDualADC_IsFinished() 和同一索引跨帧是否改变。','一帧完成后再读取；DMA 源固定、目的递增，半字搬运；生产者 DMA 与消费者算法不同时写同一帧。','只跑一次：未重新 Start；半帧数据：提前消费；随机错：连续模式无序列号/所有权。','检查 Start/IsFinished/Stop 调用顺序；连续模式使用 GetContinuousBlockSequence、GetContinuousCompletedBlock 或 GetContinuousSnapshot。','检查 CH0/CH1 配置为 fixed→block、Half Word；模块拥有 DMA_IRQHandler。','先单帧确认，后加 ping-pong 或快照；【建议新增】frame_id、buffer_owner、overrun_count 用于诊断。','参考 adc_dma_minimum、adc_ring_buffer；当前模块公开连续采集 API 见 signal_dual_adc_mspm0g3507.h。','连续 100 帧无撕裂，帧号单调，算法只读已完成帧。','改为停采—处理—重启的单帧 DMA。'),
('08','双 ADC 同步','P1','当前工程真实实现','同源两路仍有相位偏差，Lissajous 倾斜或相位随频率变化。','同步触发未确认时 STOP：不要先做相位校准。','同一信号源用等长线一分二接 PA25/PA17。','第一测量：同时查看两路 raw 的阈值上升沿样本索引差。','当前工程 ADC0/ADC1 都由 TIMG0 ZERO 事件触发；同源 baseline 应接近固定值。','恒定偏移：线缆/模拟通道延迟；漂移：触发未同步/采样率不稳；偶跳：DMA 帧错配。','检查两个 ADC 的触发发布事件、两个 DMA 通道和同一 frame 的来源。','检查输入共地、线缆等长、两路前端增益与偏置一致。','先记录 baseline，再用 Δt 校正；只对稳定、不过零歧义的信号报告相位。','当前实现：SignalDualADCPhase_Process；相位范围 [-180,180]，正值表示 Y 领先 X。','同源输入下 baseline 稳定；交换两路后符号反转。','只显示“相位无效”或用零交叉单周期法。'),
('09','测频','P1','当前工程部分实现','测频跳动、显示 0、与信号源不符。','Fs 或原始波形不可信时 STOP：不要调峰值搜索参数。','固定频率正弦/方波，先用示波器或信号源读数作真值。','第一测量：打印/显示实测 Fs、N、峰值 bin 或零交叉间隔。','时域周期法 f≈Fs/样本周期数；FFT 法 f=kFs/N。','整数倍偏差：Fs 错；低频跳：窗口内周期太少；高频跳：欠采样/阈值噪声。','确认算法使用当前帧和真实 Fs；【当前工程未集成 FFT 测频模块】不可假定存在某 API。','检查输入幅度、偏置、抗混叠与触发稳定。','先以零交叉/计数器做基线，再用 FFT；加阈值滞回和多帧平均。','参考 timer_capture_minimum、fft_minimal.c。','对 3 个已知频点误差符合比赛要求，频率不大幅跳变。','切到 Timer 捕获或零交叉测频。'),
('10','测幅、Vpp、RMS 与削顶','P1','当前工程部分实现','Vpp 偏大、RMS 不对、幅度突然变小或满屏。','出现 ADC 削顶时 STOP：不相信 Vpp/RMS，界面显示 OVER RANGE。','输入正弦，分别用小信号、中量程、接近满量程验证。','第一测量：查看 raw 的 min/max，比较是否接近 0 或 4095。','Vin=Code×Vref/FullScale；Vpp=Vmax−Vmin；Vrms=sqrt(mean(x²))（先去直流）。','Vpp 偏大：增益/标定；RMS 偏大：未去 DC；突然变小：削顶或量程切换。','确认标定系数和量程记录属于当前方案；【建议新增】adc_min/adc_max/clip_flag。','检查前端偏置、增益、输入范围与探头倍率。','先去 DC、检测削顶、使用已知幅值做两点校准。','参考 OPA-to-ADC 前端说明；当前工程 raw 为 12 bit 码值。','小/中量程误差可解释；任一端贴近 0/4095 时明确报 OVER RANGE。','显示峰峰值或相对幅度，固定增益档。'),
('11','相位、通道延迟与 baseline','P1','当前工程真实实现','相位跳 180°、同源不为 0°、低幅时乱跳。','同源 baseline 未建立、幅度不足或交叉点歧义时 STOP：不要输出“精确相位”。','同源一分二；再接 0°/90° 已知相位差源。','第一测量：同源 baseline；看两路阈值上升沿的 Δ样本与 g_phase_valid。','Δt=Δφ/(360f)；当前模块有滞回 16、最小幅度 64、幅度比 1–5 的有效性约束。','±180° 跳：周期折返/阈值错；低幅跳：噪声；随频率变：通道延迟。','检查 SignalDualADCPhase_Process 输入为同一完成帧；g_phase_degrees/g_phase_valid 是真实变量。','检查两路线缆、前端带宽、偏置与振幅比。','保存同源 baseline 并做 Δt/相位校正；无效时显示 INVALID。','当前模块：modules/signal_dual_adc_phase.c/.h。','同源稳定，接入已知相差时符号与范围正确。','退化为零交叉相位，或仅报 lead/lag。'),
('12','FFT 与频谱','P1','当前工程未实现（有库示例）','主峰 bin 错、频率轴错、幅值泄漏、频谱随机。','Fs、N、输入帧所有权或量化条件不明时 STOP：不要解释频谱。','离线固定数组或已完成 DMA 帧；先用单频正弦。','第一测量：打印 Fs、N、peak bin、输入 min/max；确认 DMA 不再写该帧。','Δf=Fs/N，f=kFs/N。例：Fs=5 MHz、N=4096，则 Δf≈1220.703 Hz；100 kHz 对应 k≈81.92，峰应在 82 附近。','整比例错：Fs/N；旁瓣大：非整周期/未加窗；主峰乱：buffer 撕裂、DC 未去除。','【当前工程未集成 FFT】先核对 fft_minimal.c 的接口与内存需求，禁止假设工程已有 FFT API。','检查输入不削顶、采样率满足奈奎斯特、前端有抗混叠。','去 DC、加窗、固定 Fs/N、只在完成帧执行；幅值标定另行验证。','库参考：MSPM0_Signal_Contest/09_examples/snippets/fft_minimal.c。','已知单频峰在理论 bin±1 内，改频后线性移动。','零交叉/Timer 测频；只显示时域波形。'),
('13','DAC 无输出','P1','当前工程未实现（有库示例）','设定电压后输出恒 0、恒满或不变化。','未在 DAC 引脚测到静态电压前 STOP：不要写 DDS/扫频。','只做固定 DC 码值输出，示波器/万用表测 DAC 引脚。','第一测量：固定中码，再测 DAC 引脚电压与使能/参考。','静态码值变化应带来单调电压变化；先确认缓冲、参考和负载。','恒 0：未使能/引脚错；恒满：量程/参考错；毛刺：更新时序/负载。','【当前工程未实现】从 dac_dc_minimum 或 dac_fixed_minimal.c 迁移，先核对目标芯片资源。','检查输出端未短路、后级输入阻抗与共地。','先固定 25%/50%/75% 三点，再接滤波与后级。','参考 dac_dc_minimum、dac_fixed_minimal.c。','三档电压单调且符合参考，接后级不失真。','外部信号源或固定 DC。'),
('14','DDS 与波形发生','P1','当前工程未实现（有库示例）','DDS 无输出、频率不对、扫频时失真或卡死。','DAC 静态输出未验证时 STOP：不能进入 DDS。','固定 1 kHz、固定幅度、固定查表长度。','第一测量：示波器看输出；打印相位累加器步进/实际更新率。','FTW≈fout×2^N/Fupdate；Fupdate 以实测定时器/DMA 更新率为准。','频率比例错：Fupdate 错；无输出：DAC/触发未起；失真：更新率不足/查表太短。','【当前工程未实现】先读 06_generator/dds/README_MINIMAL_EXAMPLE.c，不假定已有函数。','检查 DAC、低通滤波、输出幅度、示波器采样设置。','固定单频→三频→再扫频；扫频期间减少 UI 刷新。','参考 06_generator/dds/README_MINIMAL_EXAMPLE.c 与 AD9850_AD9833 文档。','1 kHz、10 kHz 等已知点稳定，频率随 FTW 单调变化。','人工切换固定频点。'),
('15','VGA、PGA、AD603 增益链','P2','当前工程未实现','幅度忽大忽小、自动量程振荡、增益不线性。','增益控制电压或供电不明确时 STOP：不要信任测幅结果。','固定输入、固定一个增益档、绕开自动控制。','第一测量：同时测输入、增益控制端、输出与 ADC raw min/max。','输出必须留出 ADC 余量，raw 不应长期贴 0/4095。','增益跳：控制环振荡；偏大：标定/接线；噪声大：带宽过宽。','【当前工程未实现】记录每档标定表，不能虚构自动量程 API。','检查供电、去耦、控制端 RC、前后级共地。','先固定档并建立标定；确认稳定后再加慢速量程切换。','参考项目模拟前端资料。','每档单调、无削顶，切换后稳定。','固定中档，显示 OVER RANGE。'),
('16','比较器、LM393 与零交叉','P2','当前工程未实现','零交叉抖动、边沿不进、频率计数错。','比较器阈值、上拉或输入摆幅未知时 STOP：不要查 MCU 捕获算法。','方波/正弦+固定阈值，先示波器看比较器输出。','第一测量：比较器输入、阈值和输出三点同时测。','输出边沿清晰且在 MCU 输入范围内；阈值应避开噪声带。','边沿抖：无滞回；恒高/低：阈值/上拉；计数错：极性/捕获边沿。','【当前工程未实现】参考 comparator_zero_cross/README.md。','检查开集输出上拉、供电、地与保护。','加滞回、设置正确边沿、统计周期而不是单次边沿。','参考 comparator_zero_cross/README.md。','输出边沿稳定，Timer 捕获频率与源一致。','用 ADC 软件零交叉并加滞回。'),
('17','运放与模拟前端','P1','当前工程部分相关','波形偏置错、振铃、削顶、ADC 数值不稳定。','前端输出未在 ADC 前实测时 STOP：不要改数字滤波。','信号源→前端→ADC，逐级断开验证。','第一测量：运放输入、输出、ADC 引脚三点的 DC 偏置和峰峰值。','ADC 输入在 0..VDDA，常用中点约 1.65 V；输出不能靠近电源轨削顶。','偏置错：参考/接法；振铃：稳定性/负载；幅度低：带宽/增益。','检查前端与 ADC 的单位、偏置去除、增益标定。','检查运放供电、共模范围、去耦、反馈网络、源阻抗。','先搭直通/固定增益，再引入滤波和可变增益。','参考 07_signal_frontend/opa_to_adc/README.md。','ADC 三档电压和正弦波形均合理，无削顶。','旁路前端或限制输入幅度。'),
('18','信号完整性','P2','当前工程相关（硬件）','高频时波形变形、地弹、串扰、偶发错误。','探头接法或地回路不可信时 STOP：不要把示波器假象当软件问题。','短地弹簧、单通道、短线、固定频率。','第一测量：ADC 引脚与信号源端对比；必要时看供电纹波。','改变探头接地/线长后现象不应根本改变。','振铃：阻抗不匹配；串扰：走线/地回流；低幅噪声：探头/屏蔽。','记录频率、探头倍率、带宽限制，避免把采样伪影当真实信号。','检查地回流、串联电阻、屏蔽、模拟数字地处理。','降低带宽/频率、缩短线、加缓冲或端接。','参考板级 SI 原则；当前工程无专用 SI API。','更换探头接法后结论一致，重复性良好。','降频、只读数不画高速波形。'),
('19','模拟滤波','P2','当前工程未实现','高频噪声混叠、幅值/相位被滤波器改变。','截止频率与采样率关系不清时 STOP：不要用滤波后数据做精密测量。','固定正弦扫频，先测无滤波基线。','第一测量：滤波器前后幅度、相位、截止附近响应。','抗混叠滤波器截止应低于 Nyquist；滤波会引入幅相误差。','噪声仍大：截止过高；幅度低：截止过低/元件误差；相位偏：群延迟。','【当前工程未实现】滤波器参数应写入硬件文档，不当作软件配置。','检查元件值、供电、布局、运放稳定性。','先确认采样率，再定截止；必要时标定频响。','参考模拟前端资料。','目标频段幅度和相位满足要求，混叠受控。','提高 Fs 或降低测量带宽。'),
('20','数字滤波','P2','当前工程未实现','滤波后延迟大、波形失真、CPU 占用高。','原始数据未正确或实时性不足时 STOP：不引入复杂滤波。','先对离线完成帧滤波，再放入实时链。','第一测量：对比滤波前后 min/max、延迟、CPU 时间。','滤波必须基于完成帧；实时链要留足采集周期。','输出滞后：群延迟；振荡：系数/溢出；卡顿：复杂度过高。','【当前工程未实现】先用库例程/离线验证，不能假定工程有某滤波 API。','检查定点位宽、饱和、buffer 边界。','先移动平均，再考虑 FIR/IIR；显示原始与滤波结果。','参考算法目录的示例。','滤波改善目标噪声且不丢帧。','关闭滤波或改为短窗口平均。'),
('21','SPI','P1','当前工程真实实现（用于 TFT）','SPI 不通、读写无响应、屏幕黑或花屏。','未看见 CS/SCK/MOSI 基本波形时 STOP：不要重写显示库。','只接一个从设备，先发固定命令/填充色。','第一测量：示波器看 CS、SCK、MOSI；确认共地。','当前 TFT 使用 SPI1；PB9 SCK、PB8 MOSI、PB6 CS。','无时钟：时钟/引脚错；有时钟无响应：CS/DC/模式错；偶发：速率/接线/SI。','检查 SPI1 的 SysConfig 与显示模块仅一处初始化。','检查供电、CS、DC、MOSI/SCK 走线与共地。','从低速、固定命令、单设备开始；确认后提高速率。','当前显示适配：signal_tft_ili9341_mspm0g3507.c。','波形正确，固定填充色稳定。','降 SPI 时钟，只显示静态文字。'),
('22','I2C','P2','当前工程未实现','扫描不到设备、NACK、总线被拉低。','SCL/SDA 未释放为高时 STOP：不要调驱动状态机。','只接一个设备和正确上拉。','第一测量：空闲时测 SCL/SDA 是否为高，起始后是否有 ACK。','空闲两线高；有合理上拉和共同电平域。','一直低：短路/设备占线；NACK：地址/供电；乱码：速率/上拉。','【当前工程未实现】先用目标器件 README/example，不能补造 API。','检查上拉电阻、地址脚、供电电压、线长。','低速扫描→单寄存器读写→再接业务。','参考器件最小例程。','稳定 ACK，重复读写一致。','取消 I2C 外设，改固定参数。'),
('23','UART','P2','当前工程未实现','没有日志、乱码、打印导致实时性下降。','串口电平/波特率不明时 STOP：不要用日志结论判断算法。','只打印固定短字符串，降低频率。','第一测量：示波器看 TX 位宽，PC 端确认端口和波特率。','固定文本无乱码；日志不能阻塞采样关键路径。','乱码：波特率/时钟；无输出：引脚/端口；卡顿：阻塞发送。','【当前工程未实现】按目标 SDK 示例配置，不假定串口对象。','检查 TX/RX 交叉、共地、电平。','用环形缓冲/低频日志；采样期关闭详细打印。','参考平台最小 UART 例程。','连续日志与采样均稳定。','屏幕显示少量状态码或 LED。'),
('24','ILI9341、OLED 与屏幕','P0','当前工程真实实现（ILI9341）','屏幕黑、花屏、刷新慢或一接屏重启。','供电掉压先 STOP；基本 SPI 波形未确认也 STOP。','当前 ILI9341 只做固定填色，不画波形。','第一测量：TFT_ILI9341_FillScreen 固定红色；再看 CS/SCK/MOSI/DC。','推荐顺序：FillScreen → Pixel → Line → Text → Waveform。当前初始化为 ROTATION_270。','全黑：供电/背光/初始化；花屏：SPI/DC/窗口；卡顿：全屏刷新/刷新频率过高。','检查 SignalTFTILI9341_MSPM0_Init(&g_tft,TFT_ILI9341_ROTATION_270) 的状态 g_tft_status。','检查 PB12 背光、PB15 DC、PB6 CS 及 SPI1 引脚。','从低速、纯色、单像素、线、文字逐层恢复；波形降采样绘制。','真实 API：TFT_ILI9341_FillScreen、DrawPixel/DrawLine/DrawString/DrawInt32。','固定填色、文字、少量线段、最终波形均按阶段通过。','关闭波形，只显示数值；关闭背光以减载。'),
('25','按键与键盘','P2','当前工程未实现','按键抖动、误触发、按下没反应。','GPIO 电平本身不可靠时 STOP：不要改界面状态机。','一个按键、轮询读取、LED 回显。','第一测量：示波器/Watch 看按下与松开 GPIO 电平。','按下/松开电平明确，去抖后一次动作只触发一次。','多次触发：无去抖；无响应：上下拉/极性；偶发：中断边沿错。','【当前工程未实现】根据实际硬件定义上下拉和引脚。','检查共地、上拉/下拉、电平兼容。','先轮询+20 ms 去抖，再添加中断。','参考 GPIO 最小例程。','连续 20 次按键无误触。','取消按键，固定参数。'),
('26','中断','P0','当前工程真实实现（DMA IRQ）','程序卡死、丢帧、ISR 不进或不退出。','ISR 高频占用/重复 owner 未解决时 STOP：不要跑复杂 UI/算法。','保留一个中断源，ISR 仅置标志或采样 GPIO。','第一测量：ISR 入口/出口翻转 GPIO，测频率和占空比。','ISR 应短、可重入风险低；当前 DMA_IRQHandler 由双 ADC 模块拥有。','不进：NVIC/映射；卡死：未清标志/死循环；丢帧：优先级/耗时过长。','检查同一 IRQ 仅一个定义，不在 main.c 重复 DMA_IRQHandler。','检查外设中断标志、优先级和电源/时钟。','把重计算移到主循环；用标志/队列交接。','当前 DMA 所有权：signal_dual_adc_mspm0g3507.c。','运行时 ISR 频率合理、无嵌套失控、帧不丢。','轮询单帧 DMA，降低刷新率。'),
('27','RAM、Flash 与 CPU','P1','当前工程真实实现','链接失败、运行随机死、加 FFT/屏幕后崩溃。','栈/堆/RAM 余量未知时 STOP：不要继续加大数组或启用 FFT。','关闭 TFT 波形、只留一帧 raw 缓冲。','第一测量：检查 map 文件 RAM/Flash，统计 g_raw_a/g_raw_b 和栈余量。','当前每个 1024×uint16_t 原始数组约 2 KiB；双路至少约 4 KiB，算法 buffer 另算。','链接溢出：静态 RAM/Flash；随机死：栈/越界；卡顿：CPU 预算不足。','检查 Debug/signal_contest_template_final.map；【建议新增】cpu_time、stack_watermark。','检查数组边界、递归、大型局部数组、DMA 目标地址。','复用 buffer、减 N、降刷新、先离线 FFT 验证。','参考 tools/ram_check/README.md 与 map 文件。','map 有余量，长时间运行无越界/复位。','N=1024、只单帧、关闭波形/FFT。'),
('28','FPGA + MCU 协同','P2','当前工程未实现','数据不同步、协议偶发错、FPGA 正常 MCU 不读。','时钟域/电平/协议未单独验证时 STOP：不要一次接完整算法链。','FPGA 固定测试字，MCU 只收一帧并显示。','第一测量：示波器/逻辑分析仪看时钟、片选、数据与有效标志。','每个接口有明确时钟域、帧边界、字序和电平。','错位：时序/边沿；偶发：CDC；全错：字序/电平/协议。','【当前工程未实现】以 FPGA 侧时序图和最小通信例程为准。','检查电平兼容、共地、上电顺序、终端。','固定 test pattern→单帧→连续流→业务算法。','参考项目 FPGA/接口文档（若选用）。','测试字无错，帧号连续。','让 MCU 独立采集/计算，FPGA 只做可选加速。'),
('29','完整数据链','P0','当前工程真实实现（采集+相位+TFT）','局部都正常，合起来错或刷新导致测量跳。','链上任一层未证实时 STOP：不得跨层调算法。','链路按“Timer→ADC→DMA→raw→处理→UI”逐层启用。','第一测量：每层只输出一个可观察量：Fs、raw 三点、frame 完成、结果、屏幕状态。','当前真实链：TIMG0→ADC0/ADC1→DMA CH0/CH1→g_raw_a/g_raw_b→SignalDualADCPhase_Process→TFT。','raw 对结果错：算法参数；结果对屏错：UI；偶发错：缓冲所有权/CPU 超时。','检查每个状态返回值与 g_adc_status、g_tft_status、g_phase_valid。','检查所有模块共地、供电、资源 owner 与刷新负载。','每次只接通一段并保留测试点；把 UI 频率降到采样之外。','当前 main.c 为真实集成参考。','连续运行且 raw、相位、屏幕状态一致。','只保留 raw 数值和相位结果，不画波形。'),
('30','精度增强','P2','当前工程未实现（有算法参考）','基础结果稳定但误差不够、频点间跳。','基线精度和 Fs 未验证时 STOP：不要先上插值/校正。','固定标准源，多频点记录误差。','第一测量：记录真值、测量值、Fs、温度/量程、误差趋势。','先校准 Fs、幅度、baseline；插值只能改善满足前提的数据。','全频偏：时基；仅边缘偏：前端频响；随机：噪声/SNR。','【当前工程未集成】参考 parabolic/Jacobsen/频响校正目录。','检查信号源精度、线缆、前端和供电稳定。','先三点标定，再使用简洁插值；保留原始结果以便回退。','参考 05_precision 目录。','误差曲线可解释，改进不损失稳定性。','使用未插值 FFT 或零交叉结果。'),
('31','猝发与触发','P2','当前工程未实现','短脉冲抓不到、触发位置漂移、预触发数据错。','触发源/时间基准不清时 STOP：不要加复杂缓存。','固定单脉冲+GPIO 触发标记，先抓一帧。','第一测量：同时看触发输入、ADC 采样、DMA 完成信号。','触发到首样点延迟应稳定可测。','漏触发：阈值/脉宽；位置漂：软件延迟；数据错：环形 buffer 所有权。','【当前工程未实现】先确定硬件触发还是软件触发，再设计 buffer。','检查电平、极性、去抖、时钟域。','固定触发→单帧→预触发环形缓冲→重复触发。','参考 ADC/DMA 环形缓冲例程。','重复触发位置稳定，帧不撕裂。','连续采样+人工保存最近一帧。'),
('32','比赛降级策略','P0','当前工程相关','高级功能反复不稳，时间不足，评分核心链仍可保。','核心“供电→采集→基本结果→展示”未稳定时 STOP：冻结所有高级功能。','建立一个可随时烧录的保底工程。','第一测量：判定哪一段已经连续稳定 10 分钟，哪一段只是偶尔成功。','保底版本应有明确版本号、参数和现场验证记录。','高级算法不稳：退基础法；UI 卡顿：只数值；双通道不稳：单通道；自动控制振荡：固定档。','记录当前工程真实 API 与冻结配置，不在最后阶段大规模重配 SysConfig。','检查备份工程、烧录线、供电和离线依赖。','采用“单帧 DMA、零交叉、固定频点、数值 UI、固定档”组合。','当前工程 main.c 与已验证模块是保底基线。','可在 5 分钟内烧录、复测、得到可信基础结果。','见速查手册保底方案总表。'),
]

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=8.5):
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if color: r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def fixed_table(doc, headers, rows, widths=None, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'; t.autofit = False
    trPr=t.rows[0]._tr.get_or_add_trPr(); th=OxmlElement('w:tblHeader'); th.set(qn('w:val'),'true'); trPr.append(th)
    for i, h in enumerate(headers):
        c=t.rows[0].cells[i]; set_cell_text(c,h,True,'FFFFFF',font_size); shade(c,BLUE)
        if widths: c.width = Cm(widths[i])
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            set_cell_text(cells[i],v,False,None,font_size)
            if widths: cells[i].width=Cm(widths[i])
            if len(t.rows)%2==0: shade(cells[i], 'F7FAFC')
    for row in t.rows:
        for c in row.cells:
            tcPr=c._tc.get_or_add_tcPr(); mar=OxmlElement('w:tcMar')
            for edge in ('top','start','bottom','end'):
                e=OxmlElement('w:'+edge); e.set(qn('w:w'),'80'); e.set(qn('w:type'),'dxa'); mar.append(e)
            tcPr.append(mar)
    return t

def configure(doc, compact=False):
    sec=doc.sections[0]; sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.7); sec.left_margin=Cm(1.75); sec.right_margin=Cm(1.75)
    normal=doc.styles['Normal']; normal.font.name='Microsoft YaHei'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); normal.font.size=Pt(9.5 if compact else 9.7)
    normal.paragraph_format.space_after=Pt(3); normal.paragraph_format.line_spacing=1.16
    for name,size,color in [('Title',25,NAVY),('Heading 1',18,NAVY),('Heading 2',13,BLUE),('Heading 3',10.5,TEAL)]:
        s=doc.styles[name]; s.font.name='Microsoft YaHei'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True
        s.paragraph_format.space_before=Pt(13 if name!='Heading 3' else 8); s.paragraph_format.space_after=Pt(5)
    if 'Callout' not in [x.name for x in doc.styles]:
        s=doc.styles.add_style('Callout', WD_STYLE_TYPE.PARAGRAPH); s.font.name='Microsoft YaHei'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); s.font.size=Pt(10); s.font.color.rgb=RGBColor.from_string(NAVY)
        s.paragraph_format.left_indent=Cm(.45); s.paragraph_format.right_indent=Cm(.45); s.paragraph_format.space_before=Pt(5); s.paragraph_format.space_after=Pt(5)
    # footer
    foot=sec.footer.paragraphs[0]; foot.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=foot.add_run('电子设计竞赛单片机 / 信号题 Debug 手册  ·  '); rr.font.size=Pt(8); rr.font.color.rgb=RGBColor.from_string('6B7280')
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); foot._p.append(fld)

def add_field(p, instr):
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), instr); p._p.append(fld)

def title(doc, main, sub):
    for _ in range(5): doc.add_paragraph()
    p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(main)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(sub); r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(TEAL)
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('基于当前 MSPM0G3507 双 ADC + DMA + ILI9341 工程整理'); r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string('4B5563')
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('版本：最终整理版 ｜ 使用前请以当前源码、.h 与 README 为准'); r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string('6B7280')
    doc.add_page_break()

def bookmark(p, anchor, bid):
    start=OxmlElement('w:bookmarkStart'); start.set(qn('w:id'),str(bid)); start.set(qn('w:name'),anchor); p._p.insert(0,start)
    end=OxmlElement('w:bookmarkEnd'); end.set(qn('w:id'),str(bid)); p._p.append(end)

def link(cell, text, anchor):
    p=cell.paragraphs[0]; p.text=''; hl=OxmlElement('w:hyperlink'); hl.set(qn('w:anchor'),anchor); hl.set(qn('w:history'),'1'); p._p.append(hl)
    r=OxmlElement('w:r'); rp=OxmlElement('w:rPr'); col=OxmlElement('w:color'); col.set(qn('w:val'),'0563C1'); rp.append(col); u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rp.append(u); r.append(rp); tx=OxmlElement('w:t'); tx.text=text; r.append(tx); hl.append(r)

def add_toc(doc, depth='1-3'):
    p=doc.add_paragraph(); p.style='Callout'; p.add_run('目录：在 Word 中右键此目录并选择“更新域”，即可生成可点击目录。')
    p=doc.add_paragraph(); add_field(p, 'TOC \\o "1-3" \\h \\z \\u')

def p(doc, text='', style=None, bold_prefix=None):
    q=doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r=q.add_run(bold_prefix); r.bold=True; q.add_run(text[len(bold_prefix):])
    else: q.add_run(text)
    return q

def label(doc, text, fill=LIGHT):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    trPr=t.rows[0]._tr.get_or_add_trPr(); th=OxmlElement('w:tblHeader'); th.set(qn('w:val'),'true'); trPr.append(th)
    c=t.cell(0,0); shade(c,fill); set_cell_text(c,text,True,NAVY,9.5); return t

def add_full_card(doc, c, bid):
    n,name,prio,status,phen,stop,minsys,first,normal,abn,software,hardware,fixes,example,accept,fallback=c
    doc.add_page_break()
    hp=doc.add_paragraph('Debug %s｜%s'%(n,name),style='Heading 1'); bookmark(hp,'dbg'+n,bid)
    fixed_table(doc,['优先级','工程状态','调试边界'],[(prio,'【%s】'%status,'先证明前一层正确，再进入本卡')],[2.1,7.0,7.1],8.5)
    sections=[
        ('现场现象',phen),('STOP 条件',stop),('最小系统',minsys),('硬件测试点',hardware),('软件测试点',software),
        ('第一步',first),('正常参考值',normal),('典型错误值 / 异常分支',abn),('SysConfig / 工程检查',software),
        ('代码检查',software),('硬件检查',hardware),('最可能原因排序','先检查供电/时钟/引脚与资源 owner；再检查本卡的软件配置；最后才评估算法与精度。'),
        ('具体解决方法',fixes),('最小 example',example),('修复后验收',accept),('仍然失败怎么办','回到本卡“最小系统”，保留第一测量结果；每次只改一个变量并记录现象变化。'),('保底方案',fallback),
    ]
    for h,txt in sections:
        doc.add_paragraph(h,style='Heading 2'); p(doc,txt)
    doc.add_paragraph('30 秒 Checklist',style='Heading 2')
    fixed_table(doc,['先看','正常时再做','不满足时'],[(first, '进入下一层或下一卡', '执行 STOP，回退最小系统')],[5.2,4.0,7.0],8)

def symptom_index(doc, quick=False):
    entries=[('电源掉压 / 重启','01'),('烧录失败 / 不启动','02'),('工程 / SysConfig 失配','03'),('模块组合失败','04'),('Timer / Fs 错','05'),('ADC 全 0 / 满量程 / 跳','06'),('DMA 只跑一次','07'),('双通道不同步','08'),('测频跳','09'),('Vpp / RMS 错','10'),('相位跳 180°','11'),('FFT 主峰错','12'),('DAC 无输出','13'),('DDS 无输出','14'),('SPI 不通','21'),('TFT 黑屏','24'),('系统卡死 / 中断','26'),('RAM 不足','27'),('完整链异常','29')]
    rows=[]
    for x,n in entries: rows.append((x,'Debug '+n,'第一个动作见对应卡顶部'))
    t=fixed_table(doc,['30 秒症状','跳转','先做什么'],rows,[5.3,2.2,8.7],8.3 if quick else 8.5)
    for row,(_,n) in zip(t.rows[1:],entries): link(row.cells[1], 'Debug '+n, 'dbg'+n)

def build_full():
    d=Document(); configure(d); title(d,'电子设计竞赛单片机 / 信号题比赛现场 Debug 手册','MSPM0 / STM32 / ADC / DMA / FFT / DDS / 显示 / 系统集成')
    d.add_heading('第 1 章｜怎么使用这本手册',0)
    label(d,'核心原则：先证明底层正确，再查算法；一次只改一个变量。')
    p(d,'本版已把旧版简卡与深度卡合并：P0/P1 内容以本 Debug Cookbook 为唯一权威入口，不再前后重复。每张卡都先给 STOP、最小系统和“第一步”。')
    fixed_table(d,['层级','先证明什么','失败时去哪'],[('L0','电源/复位/SWD','01–02'),('L1','工程与资源 owner','03–04'),('L2','Timer/ADC/DMA raw','05–08'),('L3','测量与 FFT','09–12'),('L4','生成/模拟前端','13–20'),('L5','通信与 UI','21–25'),('L6','中断与资源','26–28'),('L7','全链路','29–31'),('L8','保底交付','32')],[2.2,9.0,5.0])
    d.add_heading('第 2 章｜10 分钟故障排查流程',0)
    fixed_table(d,['时间','动作','通过标准'],[('0–1 min','断扩展模块，测 3.3 V/复位/SWD','不掉压、可下载'),('1–2 min','确认工程、.syscfg、资源 owner','干净构建、无重复 IRQ'),('2–4 min','验证 Timer Fs、ADC 三档 raw','raw 合理、Fs 实测一致'),('4–6 min','验证 DMA 完成与帧所有权','只处理已完成帧'),('6–8 min','验证算法输入/结果','Fs/N/码值可解释'),('8–10 min','最后启用 UI','FillScreen→Pixel→Line→Text→Waveform')],[2.1,8.0,6.1])
    d.add_heading('第 3 章｜症状反查',0); symptom_index(d)
    d.add_heading('第 4 章｜当前工程真实资源矩阵',0)
    label(d,'仅本表为“当前工程”事实；DAC、DDS、FFT、I2C、UART 等未集成功能须按相应卡标注处理。')
    fixed_table(d,['资源','当前占用','Owner / 注意'],RESOURCE_ROWS,[2.2,4.2,9.8])
    d.add_heading('第 5 章｜Debug Cookbook',0); add_toc(d)
    for i,c in enumerate(CARDS,1): add_full_card(d,c,i)
    d.add_page_break(); d.add_heading('附录 A｜核心公式与现场常量',0)
    fixed_table(d,['项目','公式 / 正常值','现场提醒'],[('ADC','Vin=Code×Vref/FullScale；12 bit=4095','0.5/1.65/2.5 V≈620/2048/3102'),('Timer','Fs=Ftimer/Count','当前 CPU 32 MHz，500 kSPS Count≈64，必须实测'),('FFT','Δf=Fs/N；f=kFs/N','5 MHz/4096 时 100 kHz≈bin 82'),('相位','Δt=Δφ/(360f)','先同源 baseline'),('RMS','Vrms=sqrt(mean(x²))','先去 DC，削顶则无效'),('DDS','FTW≈fout×2^N/Fupdate','Fupdate 以实测为准')],[2.0,8.2,6.0])
    d.add_heading('附录 B｜工程事实与建议变量的边界',0)
    fixed_table(d,['类别','可直接使用','不能误认为已存在'],[('真实变量','g_raw_a、g_raw_b、g_phase_degrees、g_phase_valid、g_adc_status、g_tft_status','—'),('真实 API','SignalDualADC_Init/Start/IsFinished；SignalDualADCPhase_Process；SignalTFTILI9341_MSPM0_Init','—'),('建议新增','frame_id、adc_min/adc_max、clip_flag、cpu_time、stack_watermark','须自行设计/实现并验证')],[2.3,7.7,6.2])
    d.save(OUT_FULL)

def quick_card(doc,c,bid):
    n,name,prio,status,phen,stop,minsys,first,normal,abn,software,hardware,fixes,example,accept,fallback=c
    hp=doc.add_paragraph('Debug %s｜%s'%(n,name),style='Heading 2'); bookmark(hp,'dbg'+n,bid)
    fixed_table(doc,['第一测量','STOP','正常 / 下一步'],[(first,stop,normal)],[5.1,5.4,5.7],8.0)
    p(doc,'异常：'+abn, bold_prefix='异常：')
    p(doc,'行动：'+fixes, bold_prefix='行动：')
    p(doc,'保底：'+fallback, bold_prefix='保底：')

def build_quick():
    d=Document(); configure(d,True); title(d,'电子设计竞赛单片机 / 信号题','比赛现场 Debug 速查手册')
    d.add_heading('第 1 页｜系统出问题先干什么',0)
    label(d,'禁止一上来修改 FFT。先让上一层的“可观测量”正确。')
    for x in ['电源 / 接地 / 复位','MCU / SWD / 烧录','ADC 原始码','DMA 完成帧','Raw Data 是否合理','算法：Fs / N / 标定','UI：从固定填色开始']:
        p(d,x,style='Heading 2')
    d.add_page_break(); d.add_heading('第 2–3 页｜30 秒症状索引',0); symptom_index(d,True)
    d.add_page_break(); d.add_heading('第 4 页｜当前 MSPM0 工程资源矩阵',0)
    fixed_table(d,['资源','一眼记住','Owner / 禁忌'],RESOURCE_ROWS,[2.1,4.6,9.5],8.2)
    d.add_heading('正常值（拿表就能测）',1)
    fixed_table(d,['项目','正常值','不正常先查'],[('ADC','12 bit/4095/VDDA≈3.3 V','0.5/1.65/2.5 V≈620/2048/3102'),('Timer','CPU 32 MHz；目标 500 kSPS','Count≈64，但须示波器实测 Fs'),('缓冲','N=1024；g_raw_a/g_raw_b','DMA 完成后才处理'),('相位','[-180°,180°]；Y 领先为正','先同源 baseline'),('TFT','SPI1；PB9/PB8/PB6/PB15/PB12','FillScreen 固定红色')],[2.0,7.2,7.0],8.1)
    d.add_page_break(); d.add_heading('高频 P0 / P1 卡',0); add_toc(d)
    essential=['01','02','03','04','05','06','07','08','09','10','11','12','13','14','21','24','26','27','29','32']
    by={c[0]:c for c in CARDS}
    for ii,n in enumerate(essential):
        quick_card(d,by[n],100+ii)
        if ii%2==1: d.add_page_break()
    d.add_page_break(); d.add_heading('公式：只留必须用的',0)
    fixed_table(d,['用途','公式','先决条件'],[('ADC','Vin=Code×Vref/FullScale','码值未削顶'),('Timer','Fs=Ftimer/Count','用示波器实测确认'),('FFT','Δf=Fs/N；f=kFs/N','完成帧、Fs 正确'),('相位','Δt=Δφ/(360f)','同源 baseline'),('RMS','Vrms=sqrt(mean(x²))','去 DC、未削顶'),('DDS','FTW≈fout×2^N/Fupdate','DAC 先静态通过')],[2.0,7.1,7.1],8.4)
    d.add_heading('出现这些现象，千万别做什么',1)
    fixed_table(d,['看到','不要做','立即做'],[('ADC raw 错','调 FFT','看 g_raw_a[0/512/1023]'),('供电掉压','查 SPI','断扩展模块，测 3.3 V/电流'),('DMA 未完成','跑算法','只处理完成帧'),('ADC 削顶','相信 Vpp/RMS','显示 OVER RANGE'),('TFT 黑','重写 UI','FillScreen 固定红色'),('相位跳','加高阶算法','同源 baseline'),('系统随机死','一次改五处','查 IRQ owner / RAM / 电源')],[2.5,6.1,7.6],8.2)
    d.add_page_break(); d.add_heading('保底方案总表',0)
    fixed_table(d,['正常方案','出问题','现场保底'],[('Ping-Pong DMA','不稳定','单帧 DMA'),('FFT 测频','不稳定','Timer/零交叉'),('高级插值','不稳定','三点抛物线或不插值'),('自动量程','振荡','固定档'),('TFT 波形','卡顿','只显示数值'),('双通道相位','不稳定','零交叉 / lead-lag'),('DDS 扫频','异常','手动固定频点'),('复杂滤波','CPU 不够','短窗口平均'),('自动识别','误判','输出 Unknown')],[3.7,4.0,8.5],8.3)
    d.add_heading('比赛最后一小时',1)
    fixed_table(d,['剩余时间','唯一策略'],[('>3 小时','可继续修高级功能，但随时保留可烧录基线。'),('1–3 小时','冻结核心链，功能只增不破。'),('30–60 分钟','禁止架构大改，只修明确 Bug。'),('<30 分钟','只校准、固定参数、关闭不稳定模块、切保底。'),('<10 分钟','禁止改核心代码；保存、烧录、复测、整理现场。')],[3.2,13.0],8.5)
    d.save(OUT_QUICK)

if __name__ == '__main__':
    build_full(); build_quick()
    print(OUT_FULL); print(OUT_QUICK)
