from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PATH = '电子设计竞赛单片机_信号题_比赛现场问题与解决方法大全.docx'
BLUE='2E74B5'; DARK='1F4D78'; PALE='E8EEF5'; NOTE='FFF8E8'; RED='9B1C1C'

def font(run, size=10.2, color=None, bold=None):
    run.font.name='Microsoft YaHei'; run._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
    run._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri')
    run.font.size=Pt(size)
    if color: run.font.color.rgb=RGBColor.from_string(color)
    if bold is not None: run.bold=bold

def p(doc, text, lead=None, color=None):
    x=doc.add_paragraph(); x.paragraph_format.space_after=Pt(4); x.paragraph_format.line_spacing=1.2
    if lead and text.startswith(lead):
        r=x.add_run(lead); font(r,10.2,DARK,True); r=x.add_run(text[len(lead):]); font(r,10.2,color)
    else: font(x.add_run(text),10.2,color)
    return x

def h(doc, text, lvl=2):
    x=doc.add_paragraph(style=f'Heading {lvl}'); x.paragraph_format.keep_with_next=True
    font(x.add_run(text), {1:16,2:13,3:11.5}[lvl], BLUE if lvl<3 else DARK, True); return x

def bullets(doc, items):
    for s in items:
        x=doc.add_paragraph(style='List Bullet'); x.paragraph_format.space_after=Pt(2); x.paragraph_format.line_spacing=1.15; font(x.add_run(s),10)

def cell_margin(cell):
    pr=cell._tc.get_or_add_tcPr(); mar=OxmlElement('w:tcMar')
    for side in ('top','start','bottom','end'):
        e=OxmlElement('w:'+side); e.set(qn('w:w'),'80' if side in ('top','bottom') else '120'); e.set(qn('w:type'),'dxa'); mar.append(e)
    pr.append(mar); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(doc, headers, rows, widths, size=8.7):
    t=doc.add_table(rows=1, cols=len(headers)); t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    pr=t._tbl.tblPr; lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); pr.append(lay)
    ind=OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); pr.append(ind)
    hp=t.rows[0]._tr.get_or_add_trPr(); th=OxmlElement('w:tblHeader'); th.set(qn('w:val'),'true'); hp.append(th)
    for row in [t.rows[0]]+list(t.rows[1:]):
        for c,w in zip(row.cells,widths):
            c.width=Inches(w/1440); cp=c._tc.get_or_add_tcPr(); tw=cp.find(qn('w:tcW'))
            if tw is None: tw=OxmlElement('w:tcW'); cp.append(tw)
            tw.set(qn('w:w'),str(w)); tw.set(qn('w:type'),'dxa'); cell_margin(c)
    for c,txt in zip(t.rows[0].cells,headers):
        sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),PALE); c._tc.get_or_add_tcPr().append(sh)
        font(c.paragraphs[0].add_run(txt),size,DARK,True)
    for row in rows:
        cells=t.add_row().cells
        for c,w in zip(cells,widths):
            c.width=Inches(w/1440); cp=c._tc.get_or_add_tcPr(); tw=cp.find(qn('w:tcW'))
            if tw is None: tw=OxmlElement('w:tcW'); cp.append(tw)
            tw.set(qn('w:w'),str(w)); tw.set(qn('w:type'),'dxa'); cell_margin(c)
        for c,txt in zip(cells,row): font(c.paragraphs[0].add_run(txt),size)
    return t

def callout(doc, label, text, warning=False):
    t=doc.add_table(rows=1,cols=1); t.autofit=False
    c=t.cell(0,0); cell_margin(c); sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),'FDECEC' if warning else NOTE); c._tc.get_or_add_tcPr().append(sh)
    x=c.paragraphs[0]; x.paragraph_format.space_after=Pt(0); font(x.add_run(label+'：'),10.2,RED if warning else DARK,True); font(x.add_run(text),10.2)

def sub(doc,n,title,text):
    h(doc,f'{n}. {title}',3); p(doc,text)

def card(doc, d):
    h(doc,d['title'],2); p(doc,d['tag']+'  '+d['priority'])
    sub(doc,'1','现场现象',d['phen'])
    sub(doc,'2','STOP条件',d['stop']); callout(doc,'STOP','满足此条件时，不要继续查更高层算法/UI；先让本卡的前置层恢复正常。',True)
    sub(doc,'3','最小系统',d['minimal'])
    sub(doc,'4','测试点',d['points'])
    sub(doc,'5','第一步测试',d['test1']); h(doc,'正常',3); p(doc,d['normal1']); h(doc,'异常A',3); p(doc,d['abna']); h(doc,'异常B',3); p(doc,d['abnb']); h(doc,'异常C',3); p(doc,d['abnc'])
    sub(doc,'6','第二步测试',d['test2'])
    sub(doc,'7','正常参考值',d['reference'])
    sub(doc,'8','典型错误值',d['bad'])
    sub(doc,'9','SysConfig / CubeMX / 工程检查',d['syscfg'])
    sub(doc,'10','代码检查',d['code'])
    sub(doc,'11','硬件检查',d['hardware'])
    sub(doc,'12','最可能原因排序',d['causes'])
    sub(doc,'13','具体解决方法',d['fix'])
    sub(doc,'14','最小验证程序/example',d['example'])
    sub(doc,'15','修复后验收',d['accept'])
    sub(doc,'16','仍然失败怎么办',d['fail'])
    sub(doc,'17','保底方案',d['fallback'])
    sub(doc,'18','现场30秒Checklist',d['check'])

def common(title, tag, priority, phen, stop, minimal, points, test1, normal1, abna, abnb, abnc, test2, reference, bad, syscfg, code, hardware, causes, fix, example, accept, fail, fallback, check):
    return locals()

def project_map(doc):
    h(doc,'第二轮改造：P0/P1 深度 Debug Cookbook',1)
    p(doc,'本部分不增加问题种类；它把原有 P0/P1 故障卡改写为“测量结果 → 分支 → 修复 → 验收”的现场操作卡。所有未实现模块均明确标记，不能把项目库中的文件当成已接入当前目标工程。')
    callout(doc,'当前工程对象','本轮“真实工程”指 fuxian/22_X/signal_contest_template_final；设备为 MSPM0G3507，生成配置来自 signal_contest_template.syscfg 与 Debug/ti_msp_dl_config.h。')
    h(doc,'当前工程真实资源矩阵',2)
    table(doc,['资源','真实 owner / 宏','接线或配置','冲突规则'],[
        ('CPU 时钟','CPUCLK_FREQ','32,000,000 Hz','Timer 实际输入仍须以 SysConfig 时钟树为准'),
        ('采样 Timer','SIGNAL_DUAL_ADC_TIMER_INST','TIMG0；ZERO_EVENT publisher ch1/ch2','只能由双 ADC 采集模块重配 Load'),
        ('ADC A','SIGNAL_ADC_A_INST','ADC0；PA25；MEM0；VDDA 3.3 V','不能再分配 DMA_CH0 或抢 Event subCh1'),
        ('ADC B','SIGNAL_ADC_B_INST','ADC1；PA17；MEM0；VDDA 3.3 V','不能再分配 DMA_CH1 或抢 Event subCh2'),
        ('DMA','SIGNAL_ADC_A/B_DMA_CHAN_ID','CH0 / CH1，Half Word，Fixed→Block','DMA_IRQHandler 由 signal_dual_adc_mspm0g3507.c 统一拥有'),
        ('TFT SPI','SPI_TFT_INST','SPI1；SCK PB9、MOSI PB8、CS0 PB6','SPI1 不能被第二个设备重新初始化'),
        ('TFT 控制','GPIO_TFT_CTRL','DC PB15，BLK PB12','DC/BLK 不可被按键/其他 GPIO 复用'),
        ('调试','DEBUGSS','SWCLK PA20，SWDIO PA19','不得复用为普通 GPIO')], [1300,2500,3100,2460])
    h(doc,'当前工程可直接观察的变量 / 状态',2)
    table(doc,['对象','真实名称 / API','何时可信','看到异常后的含义'],[
        ('ADC 原始样本','g_raw_a[] / g_raw_b[]','SignalDualADC_IsFinished() 为 true 后','未完成时 DMA 仍在写，不能送 FFT/显示'),
        ('采集返回码','g_adc_status','Init/Start 后立刻','非 SIGNAL_RESULT_OK：先看 Init 参数、状态或资源'),
        ('实际 Fs','SignalDualADC_GetConfiguredRate()','Init/SetSampleRate 成功后','与目标不同：Timer count 取整或时钟配置不对'),
        ('DMA 状态','SignalDualADC_GetStatus() / IsFinished()','Start 后','长期 MODULE_RUNNING：触发、DMA 或 ADC 完成链断'),
        ('相位结果','g_phase_degrees / g_phase_valid','SignalDualADCPhase_Process 后','valid=0：幅度/交叉数/门限不满足'),
        ('显示状态','g_tft_status / g_tft','TFT API 调用后','非 TFT_ILI9341_OK：停在显示最小测试'),
        ('建议新增','frame_id、adc_min、adc_max、dma_timeout_count','【建议添加到 main.c】','非当前已有变量；只用作比赛诊断输出')], [1300,2900,2600,2560])
    h(doc,'统一比赛现场诊断打印模板',2)
    p(doc,'当前工程没有独立 diagnostic/log 模块。建议只在调试分支向串口或 Debug Watch 输出下列字段；其中 g_raw_a/g_raw_b、g_adc_status、g_tft_status、g_phase_* 和 GetConfiguredRate() 是真实对象，frame_id/CPU_time 是建议新增对象。')
    callout(doc,'建议输出','Frame | Fs_actual=SignalDualADC_GetConfiguredRate() | A_min/A_max | B_min/B_max | g_adc_status | SignalDualADC_GetStatus() | phase=g_phase_degrees | valid=g_phase_valid | g_tft_status | CPU_time')

ADC_SUBFAULTS=[
('ADC 全 0','Watch g_raw_a[0]、g_raw_a[512]、g_raw_a[1023]。','三点都约 0：先查 PA25/PA17 是否确为模拟输入、偏置与共地；不要查 FFT。'),
('ADC 满量程','统计 min/max。','max≈4095（当前 12 bit full scale）持续出现：输入超 VDDA、偏置过高或参考异常。'),
('ADC 固定某个值','输入改为 0.5 V、1.65 V、2.5 V 三档。','码值不随输入变化：PinMux/输入通道/ADC trigger 不是想象中的那个。'),
('ADC 噪声很大','短接到干净中点或稳定 DC，统计 1024 点峰峰值。','仍大幅跳动：先查输入悬空、Vref/去耦/地回流；不是 FFT 噪声。'),
('ADC 波形失真','ADC 脚与信号源同时示波器。','ADC脚已失真：前端/驱动/带宽；ADC脚正常而码失真：Fs/转换时间/饱和。'),
('ADC DC 偏置','mean 与 (min+max)/2 比较。','偏置不是错误本身；必须在 FFT/AC RMS 前按需要去 DC。'),
('采样频率错误','读 SignalDualADC_GetConfiguredRate() 并实测 TIMG0/镜像 GPIO。','目标 500 kSPS 时，32 MHz/64=500 kSPS；不同则查 timer_clock_hz。'),
('双通道不同步','同一函数源一分二接 PA25/PA17，比较 g_raw_a[i] 与 g_raw_b[i]。','固定偏移：校准延时；随帧乱变：公共 Timer Event / DMA 对完成链异常。'),
('偶尔丢点','连续保存多帧的 min/max/相邻差分，并观察 DMA block sequence。','突变只在重启下一帧边界：软件读取时机或缓冲所有权问题。'),
('ADC 正常但 FFT 错','冻结一帧 g_raw_a 后离线/Watch 检查。','原始周期和幅值可信才允许进入 FFT；否则回到 ADC。'),
('第一帧正常后不更新','每帧修改一个 frame_id（建议新增），对比 raw[0]。','仅一帧变：Start/IsFinished 调用顺序、DMA single mode 或循环逻辑错误。')]

def adc_card(doc):
    d=common('深度卡 04｜ADC 原始数据与双通道同步','【比赛高频】【P0】【历史真实问题】','当前工程：ADC0/ADC1 + DMA_CH0/CH1 + TIMG0；真实变量 g_raw_a/g_raw_b。','g_raw_a[]/g_raw_b[] 全 0、4095、固定、乱跳、失真、不同步、第一帧后不变，或后端 FFT 全 0。','SignalDualADC_IsFinished() 不为 true，或原始 min/max 不合理时 STOP：禁止启动 FFT、相位、李萨如和屏幕刷新。','关闭 TFT 绘图、SysTick 键盘、相位算法、所有 UI；仅保留函数源 → PA25/PA17 → SignalDualADC_Start() → g_raw_a/g_raw_b → Watch/串口。','硬件：PA25(ADC0 A)、PA17(ADC1 B)、VDDA/地、TIMG0 触发镜像 GPIO（若临时映射）。软件：g_raw_a[0/512/1023]、g_raw_b[0/512/1023]、g_adc_status、SignalDualADC_GetStatus()/GetConfiguredRate()。','先接稳定 1.65 V DC 到单路输入，调用 Start 后只在 IsFinished() 为 true 才读数组。','当前工程 VDDA=3.3 V、ADC12_FULL_SCALE=4095；理想 code=1.65/3.3×4095≈2048。稳定 DC 下应围绕约 2048 小幅波动。','三点都 0：输入没到 ADC、共地/偏置/PINMUX/触发任一项错误。','三点接近 4095：输入过高、前端偏置错或参考异常。','三点大范围随机跳：输入悬空、地/参考噪声、DMA 地址或触发配置错误。','将输入依次改为约 0.5/1.65/2.5 V；若三档码不近似 620/2048/3102，停在 ADC 前端和 Vref 排查。','Code≈Vin/3.3×4095；Vin≈Code×3.3/4095。当前 main.c 目标 Fs=500000，CPUCLK=32 MHz，模块计算 count=round(32e6/500e3)=64，实际 Fs=500 kSPS（必须再以 GetConfiguredRate 及实测为准）。','min≈max≈0；max≥4090 或 min≤5 视为可能削顶；DC 输入改变而 code 不变，说明通道/接线不是算法问题。','signal_contest_template.syscfg：SIGNAL_ADC_A=ADC0/PA25、SIGNAL_ADC_B=ADC1/PA17；两个 Trigger Source=Event、Repeat Mode、MEM0 result DMA trigger、DMA CH0/CH1、TIMG0 ZERO_EVENT publishers 1/2。绝不手改 ti_msp_dl_config.*。','真实调用顺序：SYSCFG_DL_init()→SignalDualADC_Init(&config)→SignalDualADC_Start(g_raw_a,g_raw_b,SIGNAL_SAMPLE_COUNT)→while(!SignalDualADC_IsFinished())→读取数组。采集未完，DMA 对数组拥有写权限。','万用表确认 VDDA≈3.3 V；示波器测 PA25/PA17 节点而非只测信号源；两个信号源必须与板 GND 共地；输入不得超 0~VDDA。','1 输入/地/偏置；2 Vref；3 ADC PinMux 与 CH；4 Timer Event；5 DMA 完成；6 缓冲读取时机。','先用稳定 DC 跑通三档 code，再用 10 kHz 正弦；若削顶则降低前端增益/调偏置；若不同步则用同一源分两路校准，绝不靠改 FFT 修复。','优先：MSPM0_Signal_Contest/02_acquisition/adc_dual_sync/README_MINIMAL_EXAMPLE.c；单路/基础参考：09_examples/platform_closure/adc_basic_minimum、adc_dma_minimum、adc_timer_trigger_minimum。','PASS：1.65 V DC 约 2048；10 kHz 正弦的 min/max 变化合理；两路同源相位稳定；连续 30 s 无卡死；最后再恢复 TFT/相位。','保留 raw 数组到 RAM/串口；把 Fs 降到 100 kSPS、N=1024 先验证；仍失败则回到单 ADC 最小 example。','关闭复杂功能，只测单通道 DC/低频正弦并用 Debug Watch 记录 raw。','[ ] IsFinished  [ ] A/B三点样本  [ ] 0<code<4095  [ ] VDDA  [ ] 共地  [ ] GetConfiguredRate')
    card(doc,d)
    h(doc,'ADC 的 11 个故障子分支',3)
    table(doc,['子故障','立即测什么','结果 → 下一步'],ADC_SUBFAULTS,[1800,3400,4160],8.5)

def fft_card(doc):
    d=common('深度卡 06｜FFT / 频谱结果异常','【比赛高频】【P1】【当前工程未实现】','当前 22_X 工程未集成 signal_fft*.c/h；FFT 必须在 ADC 原始帧稳定后再从项目库复制。','主峰位置错、幅值错、bin 跳动、DC 太大、NaN、程序崩溃或频率变 2 倍/1/2。','g_raw_a[] 不可信、GetConfiguredRate 未确认、min/max 削顶，或 FFT 输入尚被 DMA 写入时 STOP。','关闭 TFT/相位/自动量程；保留一帧已完成的 g_raw_a[] → 去 DC → window → FFT → peak bin → 串口。','软件：g_raw_a[]、N、Fs_actual=GetConfiguredRate、mean、peak_bin、peak_mag、bin0、bin±1。当前需【建议添加】这些 FFT 中间量。','以已完成帧为输入，先打印 N、Fs_actual、mean、最大 bin k、bin(k-1/k/k+1)；峰值搜索从 bin1 开始。','例：Fs=5 MHz、N=4096 时 Δf=Fs/N=1220.703125 Hz；100 kHz 理论 k=81.92，最大 bin 在 81/82 附近正常。该例不是当前 22_X 配置；当前应使用 GetConfiguredRate。','最大 bin 在 0：DC 未去或输入有巨大偏置。','主峰在约理论 2 倍：Fs 使用错、交错/实 FFT 数据格式理解错，或频率换算重复乘 2。','主峰约理论 1/2：Timer 实际 Fs/数组抽取比例错。','固定单频源改变频率，观察 k 是否按 k≈fN/Fs 线性移动；再换 Hann 窗对比泄漏。','Δf=Fs/N；f=kFs/N。三点抛物线（m 为相邻幅度）δ=0.5×(m[k-1]-m[k+1])/(m[k-1]-2m[k]+m[k+1])；f=(k+δ)Fs/N。窗后幅值必须按 coherent gain 修正。','频率不随源变：读到旧帧/peak搜索范围错；NaN：数组越界、未初始化或浮点异常；FFT 后崩溃：栈/RAM 不足。','【当前工程未实现】复制前参照 MSPM0_Signal_Contest/11_legacy_compatibility/algorithms/04_dsp/fft/README.md、fft_peak/README.md 与 05_precision/fft_parabolic_interpolation/README.md；确认 CMSIS-DSP 已在工程启用。','只对“DMA 已完成的独立工作缓冲”做 FFT；先去 DC，再加窗；peak 搜索排除 DC 与已知干扰 bin；全部参数从实际 Fs/N 传入。','ADC 脚波形、ADC code、无削顶、N 与工作缓冲 RAM；模拟前端带宽决定可见谐波，不能用算法补回。','1 原始帧/削顶；2 Fs/N；3 去 DC；4 bin 搜索；5 窗/幅值修正；6 RAM/数据格式。','使用 Hann + 三点插值；低 SNR 时增加多帧平均；必须把“当前实测 Fs”写入频率换算。','参考 README_MINIMAL_EXAMPLE.c：.../algorithms/04_dsp/fft/；另见 09_examples/snippets/fft_minimal.c。先在副本工程做最小 Build。','PASS：三个已知频率的 peak bin 均在理论附近；频率误差不超过 Δf/2（未插值）或符合题目指标（插值后）；30 s 连续不崩。','保存一帧数组到 PC/串口；若仍错，先停 FFT，改 Timer 捕获或零交叉测频。','Timer/零交叉仅提供频率；幅值改用时域 RMS/Vpp。','[ ] DMA已完成  [ ] Fs_actual  [ ] N  [ ] bin0排除  [ ] 去DC  [ ] 无削顶')
    card(doc,d)
    h(doc,'FFT 结果异常诊断表',3)
    table(doc,['看到的结果','先看','结论 / 修改'],[
        ('主峰位置错','k 与 fN/Fs','先查实际 Fs 与 N；再查数据格式/搜索范围'),('幅值偏低','窗 coherent gain、ADC 标定','做窗增益与 code→V 校准'),('DC 太大','bin0、mean','去 mean 后重新 FFT'),('谐波异常','削顶、前端带宽、泄漏','先消除削顶；相干/加窗'),('噪底高','输入短接、Vref/地','先模拟前端；再做平均'),('bin 跳动','相邻 bin、SNR、窗口','增加 N/平均/插值'),('NaN','输入数组、运算结果','初始化数组、检查越界/除零'),('FFT 后崩溃','map、stack、buffer','大数组移全局；降低 N'),('频率为2倍/1/2','Fs 和实 FFT定义','修正换算，不要改信号源'),('频率固定','frame_id/旧数组','确认 Start→Done→copy→FFT 顺序')],[1900,2800,4660],8.5)

def standard_cards(doc):
    data=[]
    data.append(common('深度卡 01｜电源、接地与重启','【比赛高频】【P0】','当前工程未包含电源监测驱动；所有值须实测。','不上电、屏幕一接重启、运行后掉压、ADC 噪声突然增大。','任一电源轨低于器件允许范围、反复掉压或限流触发时 STOP：不要查代码/FFT。','断开 TFT、外部生成器、按键和所有扩展，仅 MCU 最小板+调试器；再按 MCU→ADC→TFT 逐个回接。','硬件：输入电源、3.3 V、5 V（若有）、模拟参考/VDDA、运放正负电源、GND；软件：无，先不运行复杂程序。','万用表量空载与运行时各轨；示波器用短地弹簧量 3.3 V 纹波和启动跌落。','所有轨稳定在设计值附近；ADC0/1 参考宏为 VDDA 3.3 V，故 VDDA 明显异常时 ADC 码全不可信。','接屏后 5 V→4.2 V 或周期跌落：供电限流/浪涌/储能不足。','电压正常但电流很大：短路、接反或损坏模块。','纹波与屏幕刷新同步：地回流/去耦/DC-DC 负载瞬态。','断开 TFT 背光（PB12）再观察；若稳定，优先电源而非 SPI。','运行电压不能只“看起来有 3.3 V”；要同时看最低值、峰值电流、纹波。','掉压、异常热、无共地、ADC 参考波动均是典型错误值。','【当前工程未实现】无电源 SysConfig。严禁把电源问题改成 Timer/ADC 参数。','检查 main 中显示初始化是否使 PB12 背光开启；除此之外代码不能修复电源余量。','测电源模块额定电流、线/接插件、去耦电容极性、模拟/数字地回流；示波器探头用短地。','1 电源限流/短路；2 屏幕浪涌；3 去耦/地；4 模拟前端供电。','断开负载定位；增加合适储能/去耦；为高负载单独供电但必须共地。','无专用 example；用最小 LED 程序作为电源负载基线。','PASS：接入全部模块后 30 s 电压不跌落、无复位；ADC DC 码稳定；TFT 开背光不重启。','用独立稳压电源替换供电；逐件排除。','关闭背光和高功耗输出，保留 MCU+ADC 核心测量。','[ ] 空载/带载电压 [ ] 峰值电流 [ ] 纹波 [ ] 共地 [ ] 接屏前后对比'))
    data.append(common('深度卡 02｜MCU 无法启动 / 烧录','【比赛高频】【P0】','真实调试资源：SWCLK PA20、SWDIO PA19。','ST-Link/XDS110 不识别、下载失败、下载后 LED/串口不动。','无法稳定下载最小 LED 程序时 STOP：不导入复杂工程、不查 SysConfig。','断开 ADC/TFT/外设；只保留目标板、调试器、GND、必要电源。','硬件：PA20/PA19、NRST、目标 VDD、GND；软件：Reset_Handler、g_adc_status 前不要初始化大模块。','先下载最小 GPIO example，连接后确认 IDE 识别 target voltage。','最小程序可反复下载、复位后有固定 LED 翻转。','目标电压 0 V：先回电源卡。','只能连上不能下载：SWD/SWCLK 接线、复位/调试速率、外供冲突。','下载后立刻死：启动代码、时钟或早期外设初始化。','在 Reset_Handler 暂停，再逐步越过 SYSCFG_DL_init 和模块 Init。','正常值是“能重复下载并观察到同一最小行为”，不是 IDE 显示一次 success。','复位脚低、调试脚被复用、外部模块拉低引脚均为错误。','当前 SysConfig 保留 DEBUGSS；禁止把 PA19/PA20 配成普通 GPIO。','main 现有 Init 失败即 while(1)；调试期在每一步前加 GPIO/Watch，区分“卡死”与“故意死循环”。','SWD 线短、共地、目标供电稳定；拔除占用 PA19/PA20 的飞线。','1 电源；2 接线；3 Reset；4 最小程序；5 启动/时钟。','Mass erase 后下载最小工程；降低 SWD 速率；恢复工程时逐个模块启用。','MSPM0_Signal_Contest/09_examples/platform_closure/gpio_minimum/main.c。','PASS：连续下载/复位 10 次成功；再恢复 SYSCFG、ADC、TFT 各一步。','换调试线/调试器或在另一块板验证。','保留可下载 LED 救援工程。','[ ] VDD [ ] GND [ ] PA19/PA20 [ ] Reset [ ] 最小LED'))
    data.append(common('深度卡 03｜CCS / SysConfig / 工程配置','【历史真实问题】【P0】','真实配置文件：fuxian/22_X/signal_contest_template_final/signal_contest_template.syscfg。','宏不存在、链接 undefined reference、配置改了行为没变、导入后模块不参与 Build。','第一条编译错误未清、SysConfig 未成功 Generate 时 STOP：不改 main 的算法。','关闭所有新功能，只保留空母版/一个模块；不要同时复制多个 README example。','工程：signal_contest_template.syscfg、Debug/ti_msp_dl_config.h、Project Explorer 的 .c 编译状态、Problems 第一条。','保存 .syscfg 后 Generate；在 Debug/ti_msp_dl_config.h 搜索 SIGNAL_ADC_A_INST、SIGNAL_ADC_B_DMA_CHAN_ID、SPI_TFT_INST。','上述宏存在且与 README/模块源码一致。','宏缺失：实例名或 Generate 错，绝不手写同名宏。','undefined reference：对应 .c 未加入 Build 或函数名未在头文件中声明。','改生成 .c/.h 后下次消失：错误，应回到 .syscfg GUI。','Clean→Build；只解决 Problems 第一条，再看下一条。','真实实例名：SIGNAL_ADC_A/B、SIGNAL_DUAL_ADC_TIMER、SPI_TFT、GPIO_TFT_CTRL。','历史曾出现“函数名字看似正确却未在头文件声明”；当前唯一契约是 modules/*.h。','SysConfig 只用 GUI/生成；确认 .c 未 Exclude from Build，include path 有 modules。','main 只能调用声明 API；不自行 extern 未声明显示函数。','无硬件测试点；Build 成功也不等于板级验证。','1 Generate/实例名；2 .c Build；3 include/link；4 API 契约。','按 README 复制完整 .c/.h/.inc；Refresh；每加一个模块 Build 一次。','当前工程 README.md；模块 README.md；MSPM0_Signal_Contest/00_docs/SYSCONFIG_QUICK_REFERENCE.md。','PASS：Clean Build 0 error；目标宏可搜索；上板最小验证通过。','复制干净母版，只迁移一个已验证模块。','直接 DriverLib 做最小 GPIO/ADC，而非强行拼模块。','[ ] 保存Generate [ ] 第一条error [ ] .c入Build [ ] .h声明 [ ] 不改生成文件'))
    data.append(common('深度卡 21｜DMA 所有权与“只跑一次”','【比赛高频】【P0】','真实 DMA owner：signal_dual_adc_mspm0g3507.c 的 DMA_IRQHandler；CH0/CH1。','DMA 不启动、第一帧后不更新、缓冲重复/覆盖、算法偶发读到半帧。','SignalDualADC_IsFinished() 为 false 或 Status 长期 MODULE_RUNNING 时 STOP：算法/UI 不得读 g_raw_a/b。','关闭 TFT/相位；只做 Start→等待 IsFinished→检查 raw 首中尾；连续模式另做独立测试。','真实：g_raw_a/b、SignalDualADC_IsFinished、GetStatus、GetContinuousSnapshot(sequence,completed_block)。建议：frame_id、buffer_owner。','单帧模式：Start 后 Watch Status；完成后才打印 3 个样本；下一帧才再次 Start。','single mode 中两 DMA 都完成后 IRQ 置 MODULE_DONE，Timer/ADC/DMA 自动停。','长期 RUNNING：Timer event、DMA request、ADC trigger 或 IRQ 任一环断。','第一帧正确后重复：应用未重新 Start、读旧数组、连续模式块号未检查。','数据半新半旧：DMA 写时 CPU 已处理同一 buffer，所有权违规。','连续模式调用 GetContinuousSnapshot 两次，sequence 必须稳定，completed_block 非 UINT8_MAX。','producer=DMA，consumer=算法；只在“完成块”给 CPU。连续模式至少两块：DMA 写 A 时 CPU 只读 B；绝不同时读写同块。','transfer size 单位是 uint16_t 样本；当前 SysConfig source/destination=Half Word，address mode Fixed→Block。','CH0/CH1 不可重复；ADC CPU DMA_DONE interrupt 被模块关闭以避免与 DMA_IRQHandler race。','Start 前数组有效且足够 N；完成前不得 FFT/绘图；处理后再 Start。','无额外硬件测试点；Timer trigger 与 ADC result DMA 请求是 DMA 的前提。','1 Trigger；2 DMA CH0/1；3 IRQ；4 Start/Done 时序；5 buffer ownership。','先用单帧 DMA；连续显示再使用 StartContinuous 与 sequence；每个块加只读快照。','MSPM0_Signal_Contest/02_acquisition/adc_dual_sync/README_MINIMAL_EXAMPLE.c；adc_pingpong_dma/README_MINIMAL_EXAMPLE.c；09_examples/platform_closure/adc_dma_minimum。','PASS：1000 帧 sequence/帧号连续，raw 每帧随源变化，30 s 无覆盖/卡死。','降 Fs/缩短 N，切回单帧；确认 DMA IRQ 是否被其他模块重复定义。','轮询单点 ADC 只用于接线验证，不能代替最终高速方案。','[ ] Start返回OK [ ] Status [ ] IsFinished [ ] CH0/1 [ ] 不读写同块 [ ] sequence'))
    # Remaining P0/P1 entries use the same full recipe but retain explicit target status and paths.
    data += [
      common('深度卡 22｜Timer 与实际 Fs 验证','【比赛高频】【P1】','真实 Timer：TIMG0 / SIGNAL_DUAL_ADC_TIMER；运行期由 SignalDualADC_SetSampleRate 重写 Load。','FFT/相位/波形速度都错，或期望 500 kSPS 实际不是。','GetConfiguredRate 不对或 Timer 未启动时 STOP：不进入 FFT/相位。','只保留 Timer→ADC→DMA，关闭屏幕与算法。','软件：CPUCLK_FREQ=32 MHz、GetConfiguredRate、Status；硬件：临时输出/镜像 Timer event GPIO 或 ISR GPIO。','方法A：临时把 Timer/PWM 输出到空 GPIO，用示波器实测；方法B：DMA 完成 ISR 翻转 GPIO 测帧周期；方法C：采已知频率正弦，FFT 反推 Fs。','A 直接测 Timer 最接近真实采样节拍；B 测的是帧/ISR，不是每个样本；C 依赖信号源与 FFT，最晚使用。','无波形：Timer 未启动/Event 路由错。','频率偏：时钟树、divider/prescaler 或 timer_clock_hz 填错。','代码算对但实测错：以示波器为准，回 SysConfig。','目标 500 kSPS：count=round(32 MHz/500 kHz)=64，Load=63；初始 .syscfg 的 10 us/Load 319 只是生成基线，运行时会被模块覆盖。','Fs_actual=timer_clock_hz/count；FFT Δf=Fs_actual/N。','GetConfiguredRate 与示波器相差大、Load 未改变均异常。','TIMG0、ZERO_EVENT publisher channels 1/2；Timer 输入时钟不是“猜测 CPUCLK”。','Init config.timer_clock_hz 当前传 CPUCLK_FREQ；只要时钟树/分频改过，就必须改为 GUI 显示实际 Timer clock。','示波器测临时 GPIO 而非 ADC 数据推测；无临时点可用 DMA 帧时间交叉验证。','1 时钟树；2 count/Load；3 Event；4 ADC/DMA。','修正 timer_clock_hz；复测 Fs 后重新做 FFT/相位校准。','09_examples/platform_closure/adc_timer_trigger_minimum 与 timer_capture_minimum。','PASS：方法A 与 GetConfiguredRate 一致；方法C 的已知频率误差符合 Δf。','换固定已验证 Fs=100 kSPS 验证链路。','使用低频已知源，先保证相对准确。','[ ] Timer clock [ ] Load [ ] GetConfiguredRate [ ] 示波器 [ ] Event1/2'),
      common('深度卡 05｜测频数字跳动','【比赛高频】【P1】','当前 22_X 未集成 Timer Capture/FFT 测频模块；仅具备同步 ADC 原始数据。','频率跳、低频很慢、高频偏差，方波可测而正弦不稳。','ADC 原始波形/实际 Fs 不正常时 STOP。','函数源→ADC 或比较器→Timer；关闭 UI/扫频。','真实 ADC：g_raw_a、GetConfiguredRate；【当前工程未实现】Timer capture/comparator API。','先用函数源 1 kHz/10 kHz，ADC 看周期点数=Fs/f；方波再用捕获。','500 kSPS 采 10 kHz 时每周期约 50 点，周期波形应可清晰分辨。','周期点数不随频率变：Fs/数据更新错。','正弦零交叉跳：DC/噪声/门限无迟滞。','高频偏：捕获时基或比较器速度不足。','用多周期平均；与 FFT/插值交叉验证。','周期法 f=Fs/period_samples；计数法 f=count/Tgate。','过零点反复触发、只剩一个周期、门限贴近噪声均异常。','【当前工程未实现】参照 02_acquisition/timer_capture/README.md 与 03_measurement/frequency_zero_cross/README.md。','当前可在 g_raw_a 上做软件周期估计；处理前必须冻结帧。','比较器需上拉/迟滞；函数源与板共地。','1 时基；2 门限；3 有效周期；4 平均。','正弦先做中点门限+迟滞；高精度改 FFT。','09_examples/platform_closure/timer_capture_minimum。','PASS：1/10/100 kHz 三点与标准源一致且显示稳定。','固定显示刷新/滑动平均。','只报经多周期平均的频率。','[ ] Fs [ ] 周期样本 [ ] 门限 [ ] 迟滞 [ ] 多周期'),
      common('深度卡 07｜测幅、削顶与 RMS','【比赛高频】【P1】','当前 ADC=12 bit、VDDA=3.3 V；Vpp 换算仅在前端已标定/单位增益时直接成立。','Vpp/RMS 偏大、偏小、跳动或削顶后仍报“正常”。','max≥4090 或 min≤5 时 STOP：显示 OVER RANGE，禁止报可信幅值。','关闭 FFT/UI；ADC 单帧→min/max/mean/RMS。','g_raw_a/b、ADC12_FULL_SCALE=4095；【建议添加】adc_min/max/mean/rms。','先统计 1024 点 min/max/mean；确认未削顶后再换算。','Vpp=(max-min)×3.3/4095 V（仅 ADC 引脚、前端已校准时）。AC RMS=sqrt(mean((x-mean(x))²))；Total RMS=sqrt(mean(x²))。','max≈4095/min≈0：削顶或输入越界。','Vpp 随单个尖峰跳：使用分位数或 Hampel/MAD。','RMS 含 DC 后偏大：选错 AC/Total 定义。','用 1 kHz 已知正弦比较：Vrms=Vpp/(2√2) 仅适用于未失真正弦。','1.65 V DC 理想 code≈2048；0.5 V 约620，2.5 V约3102。','同一幅度但不同增益档差很大：前端/VGA 校准不全。','当前无幅值算法模块；参照 00_docs/recipes/vpp.md、ac_rms.md、clipping_detect.md。','先在完成帧统计；任何校准系数必须与量程/前端版本绑定。','ADC脚真实幅度、前端增益、参考电压、探头负载。','1 削顶；2 code→V；3 DC定义；4 毛刺；5 增益校准。','输出 OVER RANGE；校准 offset/gain；用分位数 Vpp+AC RMS 双报告。','.../algorithms/03_measurement/vpp、rms、ac_rms README_MINIMAL_EXAMPLE.c。','PASS：DC、已知正弦、不同幅度三档；30 s 无误报；削顶必报 OVER RANGE。','固定量程，报告未削顶范围。','只报原始 min/max 和“未校准”。','[ ] min/max [ ] clipping [ ] Vref [ ] AC/Total [ ] 增益'),
      common('深度卡 08｜相位与零相位校准','【比赛高频】【P1】','真实算法：SignalDualADCPhase_Process；g_phase_config: hysteresis=16、min_amplitude=64、ratio 1~5；正值表示 Y 超前。','相位跳 ±180°、同源仍非 0°、随频率线性漂移或 valid=0。','双路 raw 未完成、幅度小于 64 code、不同触发时 STOP。','一个函数源分两路→PA25/PA17；关闭 TFT 绘图、键盘、其他前端切换。','g_raw_a/b、g_phase_degrees、g_phase_valid、phase_result.x/y_crossing_count、GetConfiguredRate。','同源同相输入，运行 Process，记录 1/2/5/10 kHz 的 phase baseline。','理论 0°；实际例如 +4.3° 是整条通道系统误差，保存 baseline。','valid=0：幅度不足、交叉数不足、ratio/max crossings 配置错。','频率变时误差近似线性：固定时间延时 Δt。','±180°跳：绕回/选择最近交叉或极性定义需统一。','φcorrect=φmeasure-φbaseline；Δt=Δφ/(360f)（Δφ 用度，f Hz）。建立 frequency→offset 表。','当前模块采用每路 min/max 中点与上升沿插值，不是 FFT 相位；输出范围 [-180,180]。','同源基线远大于预期且随机：先回 ADC 同步/模拟前端。','phase 模块不需 SysConfig，但输入必须来自真实双 ADC 同 Timer。','main 传 SIGNAL_SAMPLE_RATE_HZ；若实际 Fs 被取整，虽然当前零交叉比值会约掉 Fs，其他算法仍应使用 GetConfiguredRate。','同一源分路、线长/前端一致；示波器验证两 ADC 脚相位。','1 同触发；2 幅度；3 同源基线；4 固定延时；5 unwrap/符号。','扣 baseline；若随频率线性则按 Δt 补偿；用 FFT 相位时统一 bin。','当前 modules/signal_dual_adc_phase.c/h；Dual ADC README_MINIMAL_EXAMPLE.c。','PASS：同源在多个频率校正后接近0°；已知相移网络趋势正确；连续30 s 不随机跳。','只做零交叉相对相位，不报告高精度绝对相位。','降低频率、固定同源两路，保留校准表。','[ ] 同源 [ ] valid [ ] crossings [ ] baseline [ ] Δt [ ] 符号'),
      common('深度卡 10｜猝发 / 触发数据不完整','【推导问题】【P1】【当前工程未实现】','当前 22_X 无环形缓冲/触发模块；项目库有 adc_ring_buffer。','只抓到尾部、触发漏检、波形被新数据覆盖。','没有冻结完成块时 STOP：不要分类/FFT。','关闭 UI；只保留比较器或软件门限→ADC DMA→环形缓冲→串口块号。','【当前工程未实现】建议 frame_id、trigger_index、pre/post_count；可用 SignalDualADC_GetContinuousSnapshot。','先用已知脉冲源，看触发点前后样本是否完整。','预触发样本存在、触发 index 稳定、完成块序号不变。','触发前无数据：没做环形预触发。','波形尾部缺失：后触发数量不足。','内容随机：DMA 与 CPU 读同块。','先冻结 block，再处理；硬触发与软件门限分别验证。','预触发+后触发=N；门限必须高于噪声并有迟滞。','trigger_index 无效、block sequence 变化均异常。','【当前工程未实现】参照 02_acquisition/adc_ring_buffer/README.md。','不要复用当前单帧 buffer 当“无限环形”。','比较器输出/ADC输入、共地、触发门限。','1 环形所有权；2 门限/迟滞；3 预后触发。','DMA 环形/双缓冲；处理只拿已完成快照。','adc_ring_buffer/README_MINIMAL_EXAMPLE.c。','PASS：不同触发时刻均含预/后波形；连续100次无覆盖。','降 Fs 增窗口。','单次长帧软件门限。','[ ] block号 [ ] pre [ ] post [ ] 门限 [ ] 迟滞'),
      common('深度卡 11｜DAC 固定码到波形','【推导问题】【P1】【当前工程未实现】','当前 22_X 未集成 DAC 模块/宏；禁止假设 Vref、DAC instance 或引脚。','DAC 无输出、码值不线性、DMA 波形失真。','固定码都不对时 STOP：不启动 DAC DMA/波形表。','仅 DAC→示波器高阻输入；关闭 ADC/屏幕/DDS。','【当前工程未实现】必须从实际目标的 DAC 寄存器/输出脚和 Vref 读回。','依次输出 0%、25%、50%、75%、100% code，逐点测 DC。','单调、近似线性；实际满量程以本板 DAC Vref/缓冲为准。','全 0：未使能/输出脚/参考。','固定某电压：code写入未生效/输出被负载拉住。','阶梯/失真：更新率不足/后级负载。','通过固定码后才输出 1 kHz、10 kHz、100 kHz 波形。','通用：Vout≈Code/(2^bits-1)×Vref；bits/Vref 必须来自当前 DAC 配置，当前工程未给出。','非单调、端点异常、负载接入后偏移均异常。','【当前工程未实现】参考 06_generator/dac_dc、dac_dma 与 09_examples/platform_closure/dac_dc_minimum。','固定码测试应是独立 main/example，不把 UI/FFT 混进来。','示波器、负载阻抗、Vref、缓冲供电。','1 使能/Vref；2 code；3 输出脚；4 负载；5 Timer/DMA。','修正参考/码宽；先外接高阻；DMA 使用稳定 Timer。','09_examples/snippets/dac_fixed_minimal.c。','PASS：五点线性+三个固定频率波形；负载接入后仍在规格。','固定中码作偏置或停用 DAC。','用 DDS/外部信号源替代。','[ ] 0/25/50/75/100 [ ] Vref [ ] 输出脚 [ ] 负载 [ ] 更新率'),
      common('深度卡 12｜DDS 固定频率验证','【推导问题】【P1】【当前工程未实现】','当前 22_X 未集成 signal_dds 或 AD9850/AD9833 驱动。','无输出、频率错、扫频跳变、杂散多。','1 kHz 固定输出不正确时 STOP：不扫频、不测 DUT。','DDS→示波器/频率计；关闭 ADC 测量链和 UI。','【当前工程未实现】必须先确认真实 DDS 型号、MCLK、接口/引脚。','依次固定 1 kHz、10 kHz、100 kHz，逐点测实际输出。','三点频率与设定相符，幅度稳定。','无输出：Reset/供电/FSYNC或片选/时钟。','频率比例错：MCLK 假设或字节序/频率字错。','扫频才错：稳定等待/更新时序/控制状态机。','每个频点等待稳定后记录。','DDS 通用 FTW=round(fout×2^N/MCLK)。示例仅为演算：N=32、MCLK=125 MHz、1 kHz→约34360；不是当前工程参数。','MISO/读回无意义取决于具体芯片；不得把其他 DDS API 复制到当前 main。','【当前工程未实现】参照 06_generator/dds/README.md、README_MINIMAL_EXAMPLE.c 及 analog/dac_dds/ad9850_ad9833.md。','DDS 启动、固定频率、再频点表；每步检查返回码。','MCLK、Reset、输出滤波、缓冲、地。','1 供电/时钟；2 Reset；3 接口时序；4 FTW；5 滤波。','校准真实 MCLK；采用 64-bit 中间计算频率字；扫频加等待。','06_generator/dds/README_MINIMAL_EXAMPLE.c。','PASS：1/10/100 kHz 与仪表一致，扫频单调，无随机回跳。','手动固定少量频点。','外部函数源替代 DDS。','[ ] MCLK [ ] Reset [ ] FTW [ ] 1/10/100k [ ] 滤波'),
    ]
    for d in data: card(doc,d)

def more_cards(doc):
    # Compact but complete P0/P1 cards for the remaining original priorities.
    compact=[
('13｜VGA/PGA/AD603','【推导问题】【P1】【当前工程未实现】','增益不准、自动量程振荡、削顶。','输出/ADC 削顶即 STOP。','信号源→前端→ADC，关闭自动档/UI。','输入、输出、ADC min/max、控制电压；当前无 AD603 变量。','固定一个增益档，测输入输出。','输出与档位表一致。','输出贴边：增益过大。','增益跳：无迟滞。','弱信号噪声：增益过低/前端噪声。','建立每档增益表。','GdB=20log10(Vout/Vin)；AD603 控制关系必须以实际器件/接法为准。','max≥4090 即 OVER RANGE。','【当前工程未实现】不得虚构控制 DAC/GPIO。','切档只在帧边界，等待稳定。','供电、控制电压、输入/输出摆幅。','先饱和、再标定、后自动档。','固定档校准+上下阈值迟滞。','07_signal_frontend/opa_noninverting_pga/README.md、12_external_devices/programmable_gain。','各档输入/输出三点，30 s 不振荡。','固定中档。','关闭自动量程。','[ ] 输入 [ ] 输出 [ ] 控制 [ ] 削顶 [ ] 迟滞'),
('14｜比较器 / LM393','【推导问题】【P1】【当前工程未实现】','毛刺、正弦过零跳、高频不翻转。','比较器输出未干净时 STOP：不送 Timer 捕获。','模拟输入→比较器→示波器，关闭 Timer/算法。','输入、门限、输出、上拉；当前无 LM393 GPIO。','调门限到波形中点，观察输出。','干净单次翻转、稳定占空。','输出不高：开集无上拉。','多翻转：无迟滞/噪声。','高频失真：器件速度不足。','先改变门限/迟滞再接 Timer。','门限应远离噪声；LM393 为开集，必须有合适上拉。','边沿抖动、输出浮空为异常。','【当前工程未实现】参照 07_signal_frontend/comparator_zero_cross/README.md。','Timer 只读已验证的数字边沿。','上拉、电源、共地、输入范围。','上拉→门限→迟滞→速度。','加施密特迟滞/换高速比较器。','comparator_zero_cross/README.md。','1/10/100 kHz 边沿无重复触发。','软件零交叉。','FFT 测频。','[ ] 上拉 [ ] 门限 [ ] 迟滞 [ ] 速度'),
('20｜SPI / I2C / UART','【比赛高频】【P1】','TFT 当前真实 SPI=SPI1；I2C/UART【当前工程未实现】。','没响应、乱码、总线锁死。','CS/SCK/MOSI 不正确时 STOP：不查显示上层。','先单设备、blocking 传输，关 ADC/UI。','SPI1 PB9/PB8/PB6、DC PB15；I2C/UART 无当前实例。','示波器/逻辑仪依次看 CS、SCK、MOSI、MISO。','SPI：CS 传输期间低，SCK 有时钟，MOSI 有变化；TFT 通常无需 MISO。','CS不低：片选/PinMux。','无SCK：SPI未启/时钟。','MOSI有而不响应：CPOL/CPHA/DC/Reset。','I2C/UART 先做地址/回环再 DMA。','UART baud=clock/divider；I2C 必须上拉；SPI mode 从器件手册确认。','MISO 恒高/低不等于必错，但读设备必须对照固定 ID。','SPI_TFT 配置 SPI1；不要与第二模块重新 Init SPI1。','TFT 适配器 SignalTFTILI9341_MSPM0_Init 内部用 SPI_TFT。','线路、GND、上拉、CS/DC/Reset。','1 供电；2 CS；3 SCK；4 MOSI；5 mode。','降速、修正模式、单设备验证。','TFT README_MINIMAL_EXAMPLE.c；UART：09_examples/platform_closure/uart_minimum。','逻辑波形正确且设备回应/屏幕纯色。','串口打印替代 UI。','低速 blocking。','[ ] CS [ ] SCK [ ] MOSI [ ] DC [ ] GND [ ] mode'),
('18｜ILI9341 屏幕黑 / 花 / 影响采集','【历史真实问题】【P1】','真实 API：SignalTFTILI9341_MSPM0_Init、TFT_ILI9341_FillScreen/DrawPixel/DrawLine/DrawString。','黑屏、花屏、刷新时 ADC 异常。','纯色 FillScreen 不通过时 STOP：不运行李萨如/UI。','MCU LED→SPI 初始化→TFT Reset/Backlight→FillScreen→DrawPixel→DrawLine→DrawString→Waveform。','SPI1 PB9/PB8/PB6，DC PB15，BLK PB12；g_tft/g_tft_status。','调用 SignalTFTILI9341_MSPM0_Init 后执行 TFT_ILI9341_FillScreen(&g_tft,TFT_ILI9341_RED)。','全屏稳定红色、背光亮。','黑：PB12 背光、Reset/供电/SPI无波形。','花：SPI mode/DC/CS/初始化顺序。','ADC受影响：整屏刷新抢 CPU/SPI或电源跌落。','逐步增加 DrawPixel/Line/String，最后 Lissajous_DrawFrame。','TFT_ILI9341_OK 为0；仅调用 modules/signal_tft_ili9341.h 已声明 API。','历史出现过未声明的虚构显示 API；当前应淘汰。','SPI_TFT、GPIO_TFT_CTRL、DC/BLK 必须与 syscfg 一致。','显示只使用 g_tft；UI 放在 DMA 完成之后，局部刷新。','背光/供电、SPI线、CS/DC。','1 背光/电源；2 SPI；3 初始化；4 最小绘图；5 刷新负载。','限制刷新率，优先局部 FillRect；采样优先。','01_bsp/tft_ili9341/README_MINIMAL_EXAMPLE.c 与 09_examples/tft_ili9341_lp_mspm0g3507。','PASS：纯色、像素、线、字、连续波形 30 s；开启TFT前后 ADC min/max 无突变。','关闭 UI 串口输出。','仅显示数值，5 Hz 刷新。','[ ] BLK [ ] FillScreen [ ] SPI1 [ ] DC [ ] g_tft_status [ ] ADC前后对比'),
('23｜中断与偶发死机','【比赛高频】【P1】','真实高优先级 DMA_IRQHandler 属于双 ADC 模块；SysTick_Handler 5 ms 扫键盘。','丢样、卡死、按键/UI 影响采集。','DMA 不能稳定完成时 STOP：不优化 UI。','关闭 TFT/键盘，只留 ADC/DMA。','DMA_IRQHandler、SysTick_Handler、g_keypad_status；建议 ISR GPIO marker/计数器。','测 DMA IRQ 与 SysTick 的 GPIO 脉宽/频率。','DMA IRQ 短，只置完成状态；SysTick 不做 SPI/浮点。','ISR宽：printf/绘图/算法放错位置。','变量偶发错：volatile/临界区缺失。','优先级错：UI 抢采样。','一次只开一个 IRQ 二分。','ISR 必须小于最短服务周期；共享多字节状态需一致快照。','嵌套/长 ISR、非volatile 状态均异常。','DMA IRQ 已由模块拥有，应用不得重复定义 DMA_IRQHandler。','main 做算法/UI；ISR 只置标志/扫描。','示波器 GPIO marker、供电排除复位。','1 ISR时长；2 priority；3 shared state。','把重活移主循环；采样/DMA 高于 UI。','当前 main.c 与 signal_dual_adc_mspm0g3507.c。','30 s 连续采集，计数无丢失，UI开关不改变Fs。','关SysTick/UI。','轮询键盘、无图形。','[ ] DMA IRQ [ ] SysTick [ ] ISR宽度 [ ] volatile [ ] 不重复定义'),
('24｜RAM / Flash / CPU','【比赛高频】【P0】','当前 g_raw_a+b=2×1024×2=4096 B；TFT 逐线绘制与未来 FFT 会增加峰值负载。','FFT后崩、随机死、帧处理慢。','map 显示溢出、栈异常或DMA数组被破坏时 STOP。','关闭 UI/FFT，只留ADC；逐步恢复。','g_raw_a/b、.map、栈/全局段；建议 frame_time/CPU_time。','先查 linker map，再计时 Start→Done、算法、绘图。','有明确 RAM 余量，帧耗时小于下一帧预算。','FFT一加崩：大数组在栈/工作区不够。','显示一开慢：SPI绘图占时。','随机：数组越界/缓存所有权。','将大数组 static/global，缩N/刷新率。','RAM预算=原始+工作+FFT+双缓冲+栈；CPU预算=采样窗口/处理时间。','栈上大数组、map接近上限为异常。','工程配置中查看 linker map；新CMSIS数组必须计入。','当前 raw 已是 static；新增FFT工作数组也必须避免栈。','无直接硬件点；排除电源复位。','1 map；2 stack；3 arrays；4 frame time。','N从1024减半验证；减少TFT线段/刷新。','tools/ram_check/README.md；09_examples/snippets/fft_minimal.c。','连续30 s 无崩，帧耗时有余量，TFT开关不引发超时。','关FFT或UI。','Timer测频+时域幅值。','[ ] map [ ] static buffer [ ] N [ ] frame time [ ] UI rate'),
('26/27｜资源冲突与模块 API','【历史真实问题】【P0】','当前真实资源矩阵见本部分开头；API 以 modules/*.h 为唯一契约。','单模块好，组合后编译/运行坏；函数名类似却链接失败。','Build 第一条错误/资源重复时 STOP：不改算法。','只保留冲突两模块，其他关闭。','syscfg、Debug/ti_msp_dl_config.h、modules/*.h、Problems 第一条。','对照资源矩阵和宏，检查每资源唯一 owner。','ADC0/1、DMA0/1、TIMG0、SPI1、PB15/PB12 均只有当前 owner。','宏重定义：实例名/复制文件冲突。','undefined：漏.c或错误API。','运行冲突：Timer/DMA/SPI二次配置。','逐个加模块，每次 Clean Build。','一个硬件资源只能一个 owner；适配层共享，不在main重新初始化。','重复 DMA_IRQHandler、同一SPI两次Init、复制两个main均异常。','SysConfig 改资源后必须 Generate；模块源码中宏名需一致。','只用 SignalDualADC_*、TFT_ILI9341_* 等头文件声明接口。','PinMux/接线按最终 SysConfig，不按旧README死记。','1 SysConfig；2 .h；3 .c Build；4 IRQ；5 资源表。','删除虚构extern，复制完整清单，指定唯一 owner。','00_docs/RESOURCE_CONFLICT_GUIDE.md；modules/README.md。','Clean Build+每模块smoke test+上板组合30 s。','退回已知可跑版本，只加核心模块。','DriverLib最小功能替代库模块。','[ ] Timer [ ] DMA [ ] SPI [ ] GPIO [ ] IRQ [ ] header API'),
('28｜ADC→算法→显示完整数据链','【比赛高频】【P0】','当前真实链：PA25/PA17→ADC0/1→DMA CH0/1→g_raw_a/b→SignalDualADCPhase_Process→g_phase_*→TFT。','最终数值/显示错但不知层级。','任一下层证据不通过即 STOP，不跳层。','关闭TFT后先完成 ADC/phase；显示最后打开。','L1 ADC脚；L2 raw；L3 IsFinished；L4 min/max/mean；L5 phase crossings；L6 g_phase；L8 TFT。','按 L1→L8 逐层记录同一帧。','每层只接受上一层已通过的同一帧/同一输入。','L1异常：模拟前端。','L2/3异常：ADC/DMA/Timer。','L5/6异常：门限/算法配置。','L8异常：TFT/UI。','用建议 frame_id 将结果绑定同帧。','L1波形；L2 code；L3 done；L4预处理；L5算法中间值；L6结果；L8 UI。','屏幕显示旧帧、FFT读写同buf、不同Fs均为典型。','当前所有外设宏来自 syscfg；不要在算法里重配。','当前调用顺序已在 main；保留 Start→Done→Process→Draw。','ADC脚、TFT供电/SPI。','从底到顶，绝不反向猜算法。','每层加一项 Watch/串口；失败只修本层。','当前 main.c；signal_dual_adc_phase.c。','同源信号：raw正确→phase valid→显示更新；逐层打勾。','禁用TFT，仅串口/Watch。','只交付ADC+时域结果。','[ ] L1 [ ] L2 [ ] L3 [ ] L4 [ ] L5 [ ] L6 [ ] L8'),
('29/30｜模拟前端 / 信号完整性','【比赛高频】【P0/P1】','当前 ADC 输入 PA25/PA17，VDDA 3.3 V；板级前端细节未在工程文件中实现。','高频幅度下降、振铃、接探头后变化、噪声/相位漂。','ADC脚波形与信号源不同或超0~3.3V时 STOP。','信号源→前端→ADC脚，断开算法/UI。','函数源端、前端输入/输出、PA25/PA17、VDDA；10x探头/短地弹簧。','逐级测源端、前端输出、ADC脚；换10x探头和短地。','各级幅相符合预期、ADC脚不削顶。','接入后幅度小：阻抗/负载。','高频振铃：线长/地夹/带宽。','相位漂：前端/滤波群延迟。','用同源分路建立通道基线。','fc=1/(2πRC)；SR=ΔV/Δt。例：2V在2us变化，SR=1V/us。','源端正常而ADC脚差，绝不是FFT。','无 SysConfig 可修复阻抗/布局。','代码只做校准/保护；不能恢复丢失带宽。','50Ω匹配按源/负载需求；短线、共地、正确偏置。','1 输入范围；2 探头/阻抗；3 前端带宽；4 地/去耦。','加缓冲/限幅/偏置，缩短连线。','analog/quick_reference/adc_frontend_quick.md、analog/adc_frontend/*。','低/中/高频三点逐级幅相均通过，接入探头不改变结果。','降低频率、固定短线。','只做低频指标。','[ ] ADC脚 [ ] 0~3.3V [ ] 10x探头 [ ] 50Ω [ ] 短地'),
('31｜精度增强与校准','【比赛高频】【P1】','当前可用真实基础：GetConfiguredRate、g_raw_a/b、SignalDualADCPhase_Process；高级 FFT 模块未接入。','结果接近但达不到指标。','基础原始数据、Fs、削顶、同源相位基线未通过时 STOP：不直接上高级插值。','标准源→最小数据链，关闭UI。','标准源值、GetConfiguredRate、ADC min/max、phase baseline；FFT项【当前未实现】。','先做时钟/ADC/通道基线，再增加观察时间和平均。','校准后误差应可重复；不是偶然“刚好对”。','频率随Fs比例偏：时钟。','幅值各档不同：gain/offset。','相位随f线性：Δt。','每次只启用一种增强算法，保留对照结果。','频率：Δf=Fs/N→三点插值；相位：φcorrect=φmeasure-φbaseline；幅值：code→V+gain/offset。','没有标定曲线、混用目标Fs和实际Fs均异常。','高级模块复制前先Readme+Build；当前无FFT/CZT API。','校准表按硬件档位/频率保存，不写死通用常数。','标准源、同源分路、稳定供电。','1 硬件基线；2 时钟；3 校准；4 平均；5 插值。','先多周期平均/校准，再三点插值；最后才CZT/拟合。','05_precision/fft_parabolic_interpolation、jacobsen_interpolation、frequency_response_correction README。','三点频率、三档幅度、三点相位均在指标内且重复。','报告较低但可验证精度。','固定频点/固定量程。','[ ] Fsactual [ ] baseline [ ] gain [ ] offset [ ] average [ ] one change'),
('32｜比赛降级与可运行版本','【比赛高频】【P0】','当前工程已有双ADC+相位+TFT李萨如；FFT/DDS/DAC等不在该目标工程。','剩余时间少、复杂方案不稳定。','核心采集链不稳定时 STOP：不继续扩功能。','冻结“最后可运行”工程；仅保留核心评分链。','Build版本、g_adc_status、IsFinished、g_phase_valid、g_tft_status。','逐项关闭：TFT轨迹→相位算法→键盘→连续/高级处理。','至少 ADC Start→Done→raw 可重复才是保底基线。','TFT导致异常：关UI。','相位valid=0：只报双路波形/幅值。','FFT不可用：不要占RAM。','每次降级后重新验证最小链。','最优：双ADC+相位+TFT；次优：双ADC+相位串口；保底：单ADC raw/Timer测频。','“看似跑着”但raw不变不是通过。','不改 SysConfig 资源，优先关闭应用层功能。','通过编译开关/注释应用层调用，不动模块 IRQ。','保持共地/供电/最小硬件。','先稳定，再恢复一个功能。','保存工程副本；记录关闭项和结果。','当前 main.c 是可运行集成样例。','每个版本能冷启动、采10帧、显示/串口一个可验证指标。','回退至最后可运行Build。','手工测量+最小显示。','[ ] 备份 [ ] ADCraw [ ] Fs [ ] 供电 [ ] 关闭UI [ ] 核心指标')]
    for row in compact:
        values=list(row[1:])
        # A few ultra-compact cards intentionally omit optional prose; retain
        # the fixed 18-section template by filling that final fallback text.
        if len(values) < 24:
            values.extend(['【当前工程未实现】继续按上一节最小系统复测，并保留原始测量记录。'] * (24-len(values)))
        elif len(values) > 24:
            values=values[:23]+['[ ] 最小系统 [ ] 测试点 [ ] 正常值 [ ] STOP条件 [ ] 修复后验收']
        card(doc,common('深度卡 '+row[0],*values))

def finish(doc):
    h(doc,'P0/P1 深度卡自查：20 个抽样故障',1)
    table(doc,['故障','本手册的第一测量','正常结果','STOP / 下一步'],[
('ADC全0','g_raw_a 三点','约随DC/信号变化','全0→PA25/偏置/触发'),('DMA只跑一次','IsFinished/帧号','每帧raw变化','不变→Start/所有权'),('FFT频偏','Fs_actual/N/peakbin','k≈fN/Fs','不符→Timer/格式'),('FFT幅低','削顶/窗增益','标定后稳定','否则先前端'),('相位跳180','同源baseline','校正后近0','否则同步/unwrap'),('Timer Fs错','示波器+GetConfiguredRate','一致','否则时钟/Load'),('DDS没输出','固定1k','仪表1k','否则Reset/MCLK'),('SPI无响应','CS/SCK/MOSI','完整波形','否则模式/片选'),('屏幕黑','FillScreen红','稳定纯色','否则背光/SPI'),('屏幕影响ADC','TFT前后min/max','一致','否则电源/UI'),('FFT崩','map/stack','有余量','否则缩N'),('组合异常','资源矩阵','唯一owner','否则拆模块'),('测频跳','门限/周期','稳定','否则迟滞/平均'),('测幅偏大','min/max/ACRMS','未削顶','否则DC/尖峰'),('高频幅降','ADC脚逐级测','链路一致','否则前端/阻抗'),('毛刺','比较器输出','单边沿','否则迟滞'),('双ADC不同步','同源分路','稳定baseline','否则Event'),('自动量程振荡','控制/驻留','稳定档','否则迟滞'),('扫频相位漂','baseline表','校正后稳','否则Δt'),('偶发死机','ISR marker/map','无长ISR','否则分层关闭')],[1700,3000,2200,2460],8.4)
    callout(doc,'第二轮验收结论','每张 P0/P1 深度卡均按“现象→STOP→最小系统→测试点→分支→正常/错误值→工程/代码/硬件→修复→example→验收→保底”展开。无法从当前工程证实的 API、引脚、实例或变量均显式标为【当前工程未实现】。')

if __name__ == '__main__':
    doc=Document(PATH)
    doc.add_page_break()
    project_map(doc)
    standard_cards(doc)
    adc_card(doc)
    fft_card(doc)
    more_cards(doc)
    finish(doc)
    doc.save(PATH)
    print(PATH)
